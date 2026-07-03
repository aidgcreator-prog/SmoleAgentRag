"""
RAG Agent - smolagents + ChromaDB + Dynamic Hardware Support + Gradio UI
=======================================================================
Tabs:
  💬 General Chat  — direct LLM conversation, no retrieval
  📚 RAG Chat      — always retrieves from ChromaDB before answering
  🖼️ Vision Chat   — image upload + optional visual RAG
  🎙️ Speech to Text — transcribe audio to text using Whisper
  📂 Knowledge Base — index / manage documents
"""

import os
import gc
import warnings
import time
import threading
import base64
import re
import subprocess
import sys
import platform
import shutil
import tempfile
import contextlib
from io import BytesIO
from pathlib import Path
from typing import Optional, Callable

import gradio as gr
import torch
from PIL import Image
import base64
from byaldi import RAGMultiModalModel

# Optional: llama.cpp (GGUF) backend, alongside the HuggingFace/transformers backend.
# If llama-cpp-python isn't installed, GGUF models simply won't show up in the dropdown.
try:
    from llama_cpp import Llama as LlamaCppBackend
    LLAMA_CPP_AVAILABLE = True
except ImportError:
    LLAMA_CPP_AVAILABLE = False


def _detect_llama_cpp_gpu_support() -> bool:
    """Best-effort check of whether the installed llama-cpp-python build was
    actually compiled with GPU offload support (CUDA / Metal / ROCm).

    This is independent of torch's CUDA detection — llama-cpp-python is a
    separate compiled extension and, unless installed via SETUP.ps1 (which
    picks a hardware-matched wheel) or built from source with the right
    CMAKE_ARGS, pip installs a CPU-only build by default even on a machine
    with a working NVIDIA GPU.

    Falls back to True (optimistic) if the capability can't be introspected,
    since requesting GPU layers on a CPU-only build simply no-ops rather
    than crashing — the model still loads, just entirely on CPU.
    """
    if not LLAMA_CPP_AVAILABLE:
        return False
    try:
        import llama_cpp as _llama_cpp_module
        supports_fn = getattr(_llama_cpp_module, "llama_supports_gpu_offload", None)
        if supports_fn is not None:
            return bool(supports_fn())
    except Exception:
        pass
    return True


LLAMA_CPP_GPU_AVAILABLE = _detect_llama_cpp_gpu_support()
if LLAMA_CPP_AVAILABLE:
    print(f"[llama.cpp] GPU offload support: "
          f"{'available' if LLAMA_CPP_GPU_AVAILABLE else 'NOT available (CPU-only build — GGUF models will run on CPU)'}")

with open("image/logo.jpg", "rb") as f:
    logo_b64 = base64.b64encode(f.read()).decode()

# ──────────────────────────────────────────────────────────────────
# Localization (i18n)
# ──────────────────────────────────────────────────────────────────
LANGUAGES = {
    "kh": {
        "title": f"""
        <div style="display:flex;align-items:center;justify-content:center;gap:15px;margin-bottom:10px;">
            <img src="data:image/jpeg;base64,{logo_b64}"
                alt="LocalAiLab Logo"
                style="width:80px;height:80px;border-radius:15px;">
            <div>
                <h2 style="margin:0;">បង្កើតដោយ LocalAiLab</h2>
                <a href="https://youtube.com/@localailabkh" target="_blank">
                    📺 LocalAiLab យូធូបឆាណែល
                </a>
            </div>
        </div>
        """,
        "subtitle": "🔍 LocalAiLab · កំណែ {version}",
        "tab_general": "💬 ការសន្ទនាទូទៅ",
        "tab_general_desc": "ការសន្ទនាផ្ទាល់ជាមួយ LLM — មិនមានការទាញយកទិន្នន័យឡើយ។",
        "tab_rag": "📚 វិភាគឯកសារ",
        "tab_rag_desc": "រាល់សំណួរនឹងទាញយកទិន្នន័យពីមូលដ្ឋានចំណេះដឹងជាមុនសិន រួចទើប LLM ឆ្លើយដោយប្រើតែបរិបទនោះ។",
        "tab_vision": "🖼️​ វិភាគរូបភាព",
        "tab_vision_desc": "បង្ហោះរូបភាព និងសួរអំពីវា។",
        "tab_kb": "📂 បញ្ចូលឯកសារ",
        "label_gguf_dir": "📁 ថតម៉ូដែល GGUF (llama.cpp)",
        "gguf_dir_placeholder": "ឧទាហរណ៍៖ D:\\models\\gguf — ទុកទទេដើម្បីបិទ",
        "btn_scan_gguf": "🔍 ស្កេន",
        "gguf_scan_found": "✅ រកឃើញម៉ូដែល GGUF ចំនួន {n} ក្នុង '{dir}' — សូមមើលបញ្ជីទម្លាក់ម៉ូដែលខាងលើ",
        "gguf_scan_empty": "⚠️ រកមិនឃើញឯកសារ .gguf ក្នុង '{dir}' ទេ",
        "gguf_scan_disabled": "ℹ️ គ្មានផ្លូវត្រូវបានផ្តល់ — មុខងារម៉ូដែល GGUF ត្រូវបានបិទ (ប្រើតែម៉ូដែល HuggingFace)",
        "gguf_scan_no_backend": "⚠️ llama-cpp-python មិនទាន់បានដំឡើងទេ — មិនអាចប្រើម៉ូដែល GGUF បានទេ",
        "tab_stt": "🎙️ សំលេងទៅជាអក្សរ",
        "tab_stt_desc": "បង្ហោះ ឬថតសំឡេង ដើម្បីបំលែងវាទៅជាអក្សរ។",
        "tab_data": "📊 វិភាគទិន្នន័យ",
        "tab_data_desc": "បង្ហោះឯកសារ CSV ឬ Excel ហើយសួរ AI Agent ឱ្យវិភាគទិន្នន័យ បង្កើតក្រាហ្វិក និងរបាយការណ៍។ Agent អាចដំឡើង Python package ដែលត្រូវការដោយស្វ័យប្រវត្តិ។",
        "data_file_label": "ទម្លាក់ឯកសារ CSV / XLSX / XLS",
        "placeholder_data": "សួរអំពីទិន្នន័យរបស់អ្នក ឧ. តើនិន្នាការលក់ជារៀងរាល់ខែយ៉ាងណា?",
        "label_charts": "🖼️ ក្រាហ្វិកដែលបានបង្កើត",
        "label_report_file": "📄 ទាញយករបាយការណ៍ (Markdown)",
        "btn_reset_agent": "🔄 កំណត់ Agent ឡើងវិញ",
        "tab_about": "ℹ️ អំពីកម្មវិធី",
        "placeholder_gen": "និយាយអ្វីមួយ...",
        "placeholder_rag": "សួរអំពីឯកសាររបស់អ្នក...",
        "placeholder_vis": "សួរអំពីរូបភាព...",
        "btn_send": "ផ្ញើ ▶",
        "btn_load": "🔄 ទាញយក",
        "btn_clear": "🗑️ សម្អាត",
        "btn_index": "📥 បញ្ចូលឯកសារ",
        "btn_refresh": "🔄 បញ្ជូនឡើងវិញ",
        "btn_delete": "🗑️ លុបដែលបានជ្រើសរើស",
        "btn_clear_all": "💥 លុបទាំងអស់",
        "label_llm": "🤖 LLM",
        "label_vlm": "🎨 Vision LLM",
        "label_stt": "🎙️ ម៉ូដែលនិយាយទៅជាអក្សរ",
        "label_stt_lang": "🌐 ភាសានៃសំឡេង",
        "stt_khmer_hint": "💡 សម្រាប់ភាពត្រឹមត្រូវជាភាសាខ្មែរ សូមជ្រើសរើសម៉ូដែល 🇰🇭 ដែលបានបណ្តុះបណ្តាលជាពិសេសសម្រាប់ភាសាខ្មែរ ខាងលើ។ ម៉ូដែល Whisper ធម្មតាមានទិន្នន័យបណ្តុះបណ្តាលភាសាខ្មែរតិចតួច។",
        "btn_transcribe": "📝 បំលែងជាអក្សរ",
        "stt_audio_label": "ថត ឬបង្ហោះសំឡេង",
        "label_vis_rag": "🔍 ក៏ទាញយកបរិបទអត្ថបទផងដែរ",
        "label_vis_ret": "🖼️ អ្នកទាញយកចក្ខុវិស័យ (PDFs)",
        "label_res": "លទ្ធផល",
        "doc_table_headers": ["ប្រភព", "ប្រភេទ", "ទំព័រ", "ចំនួន Chunk"],
        "label_kb_docs": "📋 ឯកសារដែលបានបញ្ចូល",
        "header_docs": "--- \n### 📋 ឯកសារដែលបានបញ្ចូល",
        "accordion_add": "📤 បន្ថែមឯកសារ",
        "file_label": "ទម្លាក់ PDF / TXT / MD / DOCX",
        "lang_label": "ភាសា",
        "lang_options": ["Khmer", "English"],
        "status_refreshed": "🔄 បញ្ជូនឡើងវិញនូវស្ថានភាព",
        "env_fixed": "✅ បរិយាកាសត្រូវបានកែប្រែ! សូម RESTART កម្មវិធីដើម្បីអនុវត្តការផ្លាស់ប្តូរ។",
        "env_starting": "🚀 កំពុងចាប់ផ្តើមការកែប្រែបរិយាកាស...",
        "env_detected": "🔍 រកឃើញ GPU: ",
        "env_using_index": "🔗 កំពុងប្រើ PyTorch index: ",
        "env_installing": "🛠️ កំពុងដំឡើង...",
        "env_success": "✨ ការដំឡើងជោគជ័យ!",
        "env_failed": "❌ ការដំឡើងបានបរាជ័យ។ សូមពិនិត្យមើលកំណត់ត្រាខាងលើ។",
        "err_vlm": "❌ Error ក្នុងពេលបង្កើត VLM: ",
        "err_gen": "❌ Error: ",
        "err_rag": "❌ Error: ",
        "err_vis": "❌ Error: ",
        "err_upload": "⚠️ មិនអាចដោះស្រាយផ្លូវបាន: ",
        "err_unsupported": "⚠️ មិនគាំទ្រ: ",
        "err_no_files": "មិនមានឯកសារត្រូវបានបង្ហោះទេ។",
        "err_nothing_indexed": "គ្មានអ្វីត្រូវបានបញ្ចូលទេ។",
        "err_no_rows": "⚠️ មិនមានជួរណាត្រូវបានជ្រើសរើស។",
        "err_nothing_deleted": "⚠️ គ្មានអ្វីត្រូវបានលុបឡើយ។",
        "err_deleted": "🗑️ បានលុប: ",
        "err_clear_msg": "🗑️ ឯកសារទាំងអស់ត្រូវបានសម្អាត។",
        "err_empty_kb": "មូលដ្ឋានចំណេះដឹងទទេស្អាត។",
        "err_empty_visual": "empty",
        "err_visual_ready": "✅ រួចរាល់",
        "err_source_unknown": "unknown",
        "err_type_unknown": "unknown",
        "err_page_empty": "—",
        "err_status_bar": "📚 Text chunks: {n}  |  Visual index: {vis}  |  Device: {dev}",
        "err_reasoning": "🧠 ការគិត (ចុចដើម្បីពង្រីក)",
        "err_answer": "💬 ចម្លើយ",
        "err_src_none": " | គ្មានឯកសារបញ្ចូលឡើយ",
        "err_src_list": " | ប្រភព: ",
        "err_elapsed": "⏱ {elapsed:.1f}s | model: <code>{model}</code> ({dev})"
    },
    "en": {
        "title": "🔍 RAG Agent",
        "subtitle": "smolagents · ChromaDB · Gemma4 / Qwen3.6 · {device} · v{version} · Developed by LocalAiLab",
        "tab_general": "💬 General Chat",
        "tab_general_desc": "Direct conversation with the LLM — no retrieval.",
        "tab_rag": "📚 RAG Chat",
        "tab_rag_desc": "Every question retrieves from the knowledge base first, then the LLM answers using only that context.",
        "tab_vision": "🖼️ Vision Chat",
        "tab_vision_desc": "Upload an image and ask questions about it.",
        "tab_kb": "📂 Knowledge Base",
        "label_gguf_dir": "📁 GGUF Model Folder (llama.cpp)",
        "gguf_dir_placeholder": "e.g. D:\\models\\gguf — leave empty to disable",
        "btn_scan_gguf": "🔍 Scan",
        "gguf_scan_found": "✅ Found {n} GGUF model(s) in '{dir}' — check the model dropdowns above",
        "gguf_scan_empty": "⚠️ No .gguf files found in '{dir}'",
        "gguf_scan_disabled": "ℹ️ No path provided — GGUF models disabled (HuggingFace models only)",
        "gguf_scan_no_backend": "⚠️ llama-cpp-python is not installed — GGUF models unavailable",
        "tab_stt": "🎙️ Speech to Text",
        "tab_stt_desc": "Upload or record audio to transcribe it into text.",
        "tab_data": "📊 Data Analysis",
        "tab_data_desc": "Upload a CSV or Excel file and ask the AI agent to analyze it, build charts, and write a report. The agent can install any Python packages it needs.",
        "data_file_label": "Drop CSV / XLSX / XLS",
        "placeholder_data": "Ask about your data, e.g. what's the monthly sales trend?",
        "label_charts": "🖼️ Generated Charts",
        "label_report_file": "📄 Download Report (Markdown)",
        "btn_reset_agent": "🔄 Reset Agent",
        "tab_about": "ℹ️ About",
        "placeholder_gen": "Say anything ...",
        "placeholder_rag": "Ask about your documents ...",
        "placeholder_vis": "Ask about the image ...",
        "btn_send": "Send ▶",
        "btn_load": "🔄 Load",
        "btn_clear": "🗑️ Clear",
        "btn_index": "📥 Index files",
        "btn_refresh": "🔄 Refresh",
        "btn_delete": "🗑️ Delete Selected",
        "btn_clear_all": "💥 Clear ALL",
        "label_llm": "🤖 LLM",
        "label_vlm": "🎨 Vision LLM",
        "label_stt": "🎙️ Speech-to-Text Model",
        "label_stt_lang": "🌐 Audio Language",
        "stt_khmer_hint": "💡 For better Khmer accuracy, choose one of the 🇰🇭 Khmer-tuned models above. Vanilla Whisper models only saw a small amount of Khmer during training.",
        "btn_transcribe": "📝 Transcribe",
        "stt_audio_label": "Record or upload audio",
        "label_vis_rag": "🔍 Also retrieve text context",
        "label_vis_ret": "🖼️ Visual retriever (PDFs)",
        "label_res": "Result",
        "doc_table_headers": ["Source", "Type", "Pages", "Chunks"],
        "label_kb_docs": "📋 Indexed Documents",
        "header_docs": "--- \n### 📋 Indexed Documents",
        "accordion_add": "📤 Add Documents",
        "file_label": "Drop PDF / TXT / MD / DOCX",
        "lang_label": "Language",
        "lang_options": ["English", "Khmer"],
        "status_refreshed": "🔄 Status Refreshed",
        "env_fixed": "✅ Environment fixed! Please RESTART the app to apply changes.",
        "env_starting": "🚀 Starting environment fix...",
        "env_detected": "🔍 Detected GPU: ",
        "env_using_index": "🔗 Using PyTorch index: ",
        "env_installing": "🛠️ Installing...",
        "env_success": "✨ Installation successful!",
        "env_failed": "❌ Installation failed. Check the logs above.",
        "err_vlm": "❌ Error during VLM generation: ",
        "err_gen": "❌ Error: ",
        "err_rag": "❌ Error: ",
        "err_vis": "❌ Error: ",
        "err_upload": "⚠️ Cannot resolve path: ",
        "err_unsupported": "⚠️ Unsupported: ",
        "err_no_files": "No files uploaded.",
        "err_nothing_indexed": "Nothing indexed.",
        "err_no_rows": "⚠️ No rows selected.",
        "err_nothing_deleted": "⚠️ Nothing deleted.",
        "err_deleted": "🗑️ Deleted: ",
        "err_clear_msg": "🗑️ All documents cleared.",
        "err_empty_kb": "The knowledge base is empty.",
        "err_empty_visual": "empty",
        "err_visual_ready": "✅ ready",
        "err_source_unknown": "unknown",
        "err_type_unknown": "unknown",
        "err_page_empty": "—",
        "err_status_bar": "📚 Text chunks: {n}  |  Visual index: {vis}  |  Device: {dev}",
        "err_reasoning": "🧠 Reasoning (click to expand)",
        "err_answer": "💬 Answer",
        "err_src_none": " | no docs indexed",
        "err_src_list": " | sources: ",
        "err_elapsed": "⏱ {elapsed:.1f}s | model: <code>{model}</code> ({dev})"
    }
}

class HardwareManager:
    @staticmethod
    def detect_nvidia_cuda_version() -> Optional[str]:
        try:
            import subprocess
            result = subprocess.run(["nvidia-smi"], capture_output=True, text=True, check=False)
            if result.returncode == 0:
                for line in result.stdout.splitlines():
                    if "CUDA Version:" in line:
                        return line.split("CUDA Version:")[1].strip()
            return None
        except Exception:
            return None

    @staticmethod
    def get_required_torch_index(cuda_version: str) -> str:
        if not cuda_version:
            return "https://download.pytorch.org/whl/cpu"

        try:
            v_parts = cuda_version.split('.')
            major = v_parts[0]
            minor = v_parts[1]

            if major == "11":
                return "https://download.pytorch.org/whl/cu118"
            elif major == "12":
                if minor in ("1", "2", "3"):
                    return "https://download.pytorch.org/whl/cu121"
                elif minor in ("4", "5", "6"):
                    return "https://download.pytorch.org/whl/cu124"
                elif minor in ("7", "8", "9"):
                    return "https://download.pytorch.org/whl/cu128"
                else:
                    return "https://download.pytorch.org/whl/cu128"
            return "https://download.pytorch.org/whl/cpu"
        except Exception:
            return "https://download.pytorch.org/whl/cpu"

    @staticmethod
    def get_system_status() -> dict:
        status = {
            "gpu_brand": "none",
            "cuda_version": None,
            "torch_cuda_available": False,
            "current_device": "cpu",
            "recommended_index": "https://download.pytorch.org/whl/cpu"
        }

        cuda_ver = HardwareManager.detect_nvidia_cuda_version()
        if cuda_ver:
            status["gpu_brand"] = "nvidia"
            status["cuda_version"] = cuda_ver
            status["recommended_index"] = HardwareManager.get_required_torch_index(cuda_ver)

        try:
            import torch
            status["torch_cuda_available"] = torch.cuda.is_available()
            if status["torch_cuda_available"]:
                status["current_device"] = "cuda"
            elif hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
                status["current_device"] = "mps"
            else:
                status["current_device"] = "cpu"
        except Exception:
            pass

        if status["gpu_brand"] == "none":
            try:
                import subprocess
                result = subprocess.run(["wmic", "path", "win32_VideoController", "get", "name"],
                                        capture_output=True, text=True, check=False)
                if "Radeon" in result.stdout or "AMD" in result.stdout:
                    status["gpu_brand"] = "amd"
                    status["recommended_index"] = "https://download.pytorch.org/whl/rocm6.2"
            except Exception:
                pass

        return status

    @staticmethod
    def fix_environment_gen():
        status = HardwareManager.get_system_status()
        if status["torch_cuda_available"]:
            yield "✅ Environment is already correctly configured for GPU!", "cuda", ""
            return

        index = status["recommended_index"]
        yield "🚀 Starting environment fix...", "cuda", "🚀 Starting environment fix..."
        yield f"🔍 Detected GPU: {status['gpu_brand'].upper()}", "cuda", f"🔍 Detected GPU: {status['gpu_brand'].upper()}"
        yield f"🔗 Using PyTorch index: {index}", "cuda", f"🔗 Using PyTorch index: {index}"

        try:
            cmd = [sys.executable, "-m", "pip", "install", "torch", "torchvision", "torchaudio", "--index-url", index]
            yield f"Running: {' '.join(cmd)}", "cuda", f"Running: {' '.join(cmd)}"

            process = subprocess.Popen(cmd, stdout=subprocess.PIPE, stderr=subprocess.STDOUT, text=True, bufsize=1)

            for line in process.stdout:
                yield "🛠️ Installing...", "cuda", line.strip()

            process.wait()

            if process.returncode == 0:
                yield "✨ Installation successful!", "cuda", "✨ Installation successful!"
                yield "✅ Environment fixed! Please RESTART the app to apply changes.", "cuda", "✅ Environment fixed! Please RESTART the app to apply changes."
            else:
                yield f"❌ Installation failed with exit code {process.returncode}", "cpu", f"❌ Installation failed with exit code {process.returncode}"
                yield "❌ Installation failed. Check the logs above.", "cpu", "❌ Installation failed. Check the logs above."

        except Exception as e:
            yield f"❌ Error: {str(e)}", "cpu", f"❌ Error: {str(e)}"


# ──────────────────────────────────────────────────────────────────
# llama.cpp (GGUF) model backend
# ──────────────────────────────────────────────────────────────────
# The GGUF model folder is NOT hardcoded. It defaults to the
# LLAMA_CPP_MODEL_DIR environment variable (empty/unset = feature off),
# and can also be set or changed at runtime from the "📁 GGUF Model Folder"
# box in the UI (see rescan_gguf_models()) — no code edits or restart
# required either way.
LLAMA_CPP_MODEL_DIR = os.environ.get("LLAMA_CPP_MODEL_DIR", "").strip()


def discover_gguf_models(folder: Optional[str] = None) -> dict:
    """Scan a folder for .gguf files and return {label: path} entries
    that can be merged straight into MODEL_OPTIONS.

    If `folder` is not given, the current global LLAMA_CPP_MODEL_DIR is
    used. An empty/unset folder simply yields no GGUF models — the app
    works fine without one.
    """
    if not LLAMA_CPP_AVAILABLE:
        return {}
    folder = folder if folder is not None else LLAMA_CPP_MODEL_DIR
    if not folder:
        return {}
    p = Path(folder)
    if not p.exists():
        return {}
    mode_tag = "llama.cpp | GPU" if LLAMA_CPP_GPU_AVAILABLE else "llama.cpp | CPU only — slow"
    found = {}
    for f in sorted(p.glob("*.gguf")):
        size_gb = f.stat().st_size / (1024 ** 3)
        label = f"🦙 {f.stem}  (~{size_gb:.1f} GB | {mode_tag})"
        found[label] = str(f)
    return found


@contextlib.contextmanager
def _capture_native_output():
    """Capture stdout/stderr at the OS file-descriptor level.

    llama.cpp is a C/C++ library — its logging (even when llama-cpp-python's
    `verbose=True` is set) writes directly to the process's real fd 1/2,
    bypassing Python's `sys.stdout`/`sys.stderr` entirely. A normal
    `contextlib.redirect_stdout` can't see it. Redirecting the actual file
    descriptors is the only way to capture the real ggml/llama.cpp error
    line (bad magic number, unknown architecture, missing split-file part,
    truncated download, etc.) instead of just the generic Python wrapper
    exception.
    """
    stdout_fd = sys.stdout.fileno()
    stderr_fd = sys.stderr.fileno()
    saved_stdout_fd = os.dup(stdout_fd)
    saved_stderr_fd = os.dup(stderr_fd)
    tmp = tempfile.TemporaryFile(mode="w+b")
    try:
        sys.stdout.flush()
        sys.stderr.flush()
        os.dup2(tmp.fileno(), stdout_fd)
        os.dup2(tmp.fileno(), stderr_fd)
        yield tmp
    finally:
        sys.stdout.flush()
        sys.stderr.flush()
        os.dup2(saved_stdout_fd, stdout_fd)
        os.dup2(saved_stderr_fd, stderr_fd)
        os.close(saved_stdout_fd)
        os.close(saved_stderr_fd)


def _raise_gguf_load_error(model_path: str, n_ctx: int, n_gpu_layers: int,
                            original_exc: Exception):
    """Reload the model once more with verbose=True, capturing llama.cpp's
    native log output, and raise a RuntimeError that includes the *real*
    reason the load failed — instead of just the generic wrapper message
    that verbose=False normally leaves us with.
    """
    diag_text = ""
    try:
        with _capture_native_output() as tmp:
            try:
                LlamaCppBackend(
                    model_path=model_path,
                    n_ctx=n_ctx,
                    n_gpu_layers=n_gpu_layers,
                    flash_attn=False,
                    verbose=True,
                )
            except Exception:
                pass  # we only care about the captured log, not this exception
        tmp.seek(0)
        diag_text = tmp.read().decode("utf-8", errors="replace").strip()
        tmp.close()
    except Exception:
        diag_text = ""

    diag_tail = "\n".join(diag_text.splitlines()[-15:]) if diag_text else \
        "(no additional native log captured — this environment may not " \
        "support fd-level output redirection)"

    raise RuntimeError(
        f"Failed to load GGUF model '{model_path}'.\n\n"
        f"Real llama.cpp log from a diagnostic reload:\n"
        f"{'-' * 60}\n{diag_tail}\n{'-' * 60}\n\n"
        "Common causes:\n"
        "  1. The file is corrupted or an incomplete download — check its "
        "size against the source repo.\n"
        "  2. It's one part of a split multi-file GGUF (e.g. "
        "'-00001-of-00002.gguf') and the other part(s) are missing from "
        "the folder — keep all parts together and point at part 1.\n"
        "  3. The installed llama-cpp-python build's llama.cpp version is "
        "too old for this model's architecture/metadata (common for newer "
        "or uncommon MoE/hybrid-attention checkpoints) — try `pip install "
        "llama-cpp-python --upgrade --force-reinstall` (or rerun "
        "SETUP.bat), or re-download the file.\n\n"
        f"Original error: {original_exc}"
    ) from original_exc


class LlamaCppModel:
    """Minimal smolagents-compatible Model wrapper around llama-cpp-python.

    Satisfies the smolagents Model contract: callable/`generate()` taking a
    list of chat messages and returning an object with a `.content` attribute.
    Lets .gguf models sit in the same MODEL_OPTIONS dropdown, and be passed
    to smolagents' CodeAgent, as the HuggingFace TransformersModel entries.
    """

    def __init__(self, model_path: str, n_ctx: int = 16384,
                 n_gpu_layers: int = -1, temperature: float = 0.6,
                 top_p: float = 0.95, max_new_tokens: int = 512,
                 flash_attn: bool = True):
        self.model_id = model_path
        self.model_path = model_path
        self.temperature = temperature
        self.top_p = top_p
        self.max_new_tokens = max_new_tokens
        print(f"[llama.cpp] Loading '{model_path}' (n_gpu_layers={n_gpu_layers}, flash_attn={flash_attn}) …")
        try:
            self.llm = LlamaCppBackend(
                model_path=model_path,
                n_ctx=n_ctx,
                n_gpu_layers=n_gpu_layers,   # -1 = offload all layers if the build supports GPU
                flash_attn=flash_attn,
                verbose=False,
            )
        except TypeError:
            # Installed llama-cpp-python is too old to accept flash_attn= at all
            print("[llama.cpp] This llama-cpp-python build doesn't support the "
                  "flash_attn parameter — loading without it.")
            self.llm = LlamaCppBackend(
                model_path=model_path,
                n_ctx=n_ctx,
                n_gpu_layers=n_gpu_layers,
                verbose=False,
            )
        except Exception as e:
            if flash_attn:
                # Some architectures (e.g. hybrid local/global attention models
                # with mismatched per-layer head dims) reject flash attention.
                # Retry without it — but if THIS also fails, the problem isn't
                # flash attention at all (e.g. an unsupported/newer GGUF
                # architecture, a corrupted download, or a missing split-file
                # part), so run one more diagnostic reload with verbose
                # logging captured, and surface the real llama.cpp error
                # instead of a confusing raw double traceback.
                print(f"[llama.cpp] flash_attn=True failed to load ({e}); "
                      f"retrying without flash attention …")
                try:
                    self.llm = LlamaCppBackend(
                        model_path=model_path,
                        n_ctx=n_ctx,
                        n_gpu_layers=n_gpu_layers,
                        flash_attn=False,
                        verbose=False,
                    )
                except Exception as e2:
                    _raise_gguf_load_error(model_path, n_ctx, n_gpu_layers, e2)
            else:
                # Loading without flash attention failed on the first try —
                # still worth a diagnostic reload to surface the real reason
                # (corrupted file, missing split part, unsupported arch, …)
                # rather than the bare wrapper exception.
                _raise_gguf_load_error(model_path, n_ctx, n_gpu_layers, e)

    @staticmethod
    def _flatten_content(content) -> str:
        if isinstance(content, list):
            return " ".join(
                c.get("text", "") if isinstance(c, dict) else str(c) for c in content
            )
        return str(content)

    def _to_plain_messages(self, messages: list) -> list:
        plain = []
        for m in messages:
            role = getattr(m, "role", None) or m.get("role")
            content = getattr(m, "content", None)
            if content is None:
                content = m.get("content")
            plain.append({"role": str(role), "content": self._flatten_content(content)})
        return plain

    def generate(self, messages: list, stop_sequences: Optional[list] = None, **kwargs):
        from smolagents.models import ChatMessage
        plain_messages = self._to_plain_messages(messages)
        out = self.llm.create_chat_completion(
            messages=plain_messages,
            stop=stop_sequences or [],
            max_tokens=kwargs.get("max_tokens", self.max_new_tokens),
            temperature=kwargs.get("temperature", self.temperature),
            top_p=kwargs.get("top_p", self.top_p),
        )
        text = out["choices"][0]["message"]["content"]
        return ChatMessage(role="assistant", content=text)

    # smolagents Model instances are called directly in some code paths
    def __call__(self, messages: list, stop_sequences: Optional[list] = None, **kwargs):
        return self.generate(messages, stop_sequences=stop_sequences, **kwargs)


if torch.cuda.is_available():
    DEVICE = "cuda"
elif hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
    DEVICE = "mps"
else:
    DEVICE = "cpu"

TORCH_DTYPE = torch.float16 if DEVICE != "cpu" else torch.float32

warnings.filterwarnings("ignore", category=UserWarning, module="torch")

# ──────────────────────────────────────────────────────────────────
# Developer branding (LocalAiLab logo, shown in the header)
# ──────────────────────────────────────────────────────────────────
_LOGO_PATH = Path(__file__).parent / "assets" / "logo.jpg"

def _load_logo_b64() -> str:
    try:
        data = _LOGO_PATH.read_bytes()
        return base64.b64encode(data).decode("utf-8")
    except Exception:
        return ""

DEVELOPER_LOGO_B64 = _load_logo_b64()
DEVELOPER_NAME = "LocalAiLab"
APP_VERSION = "1.2.0-beta"

# ──────────────────────────────────────────────────────────────────
# About tab — always shown bilingually (Khmer first, English below),
# independent of the language dropdown. Kept as static strings so it
# stays accurate to the current tab order / feature set at a glance.
# ──────────────────────────────────────────────────────────────────
def _about_content_kh(device: str, version: str) -> str:
    return f"""
### 🔖 កំណែ {version}

ភ្នាក់ងារ RAG មូលដ្ឋាន ពហុភាសា (ខ្មែរ/អង់គ្លេស) និងពហុម៉ូដាល — សន្ទនាទូទៅ វិភាគឯកសារតាមរយៈ RAG យល់ដឹងរូបភាព
បំលែងសំឡេងទៅជាអក្សរ វិភាគទិន្នន័យ CSV/Excel ដោយ AI Agent និងគ្រប់គ្រងមូលដ្ឋានចំណេះដឹង — ដំណើរការទាំងស្រុងនៅលើកុំព្យូទ័ររបស់អ្នក
ដោយប្រើម៉ូដែល HuggingFace (transformers) ឬម៉ូដែលមូលដ្ឋាន GGUF តាមរយៈ llama.cpp។

---

## ផ្ទាំង

| ផ្ទាំង | ការពិពណ៌នា |
|---|---|
| 💬 ការសន្ទនាទូទៅ | ការសន្ទនាផ្ទាល់ជាមួយ LLM — មិនមានការទាញយក |
| 📚 ការសន្ទនា RAG | ទាញយកពីមូលដ្ឋានចំណេះដឹងជាមុន រួចឆ្លើយ |
| 🖼️ ការសន្ទនាចក្ខុវិស័យ | យល់ដឹងរូបភាព ជាមួយបរិបទអត្ថបទ |
| 🎙️ និយាយទៅជាអក្សរ | បំលែងសំឡេងជាអក្សរ ដោយប្រើ Whisper |
| 📊 វិភាគទិន្នន័យ | Agent វិភាគ CSV/Excel បង្កើតក្រាហ្វិក និងរបាយការណ៍ |
| 📂 មូលដ្ឋានចំណេះដឹង | បង្ហោះ (PDF / TXT / MD / DOCX) និងគ្រប់គ្រងឯកសារ |

## ស្ថាបត្យកម្ម

| សមាសធាតុ | លម្អិត |
|---|---|
| LLM | ម៉ូដែល HuggingFace (Qwen3, Gemma-4) **ឬ** ម៉ូដែល GGUF មូលដ្ឋានតាមរយៈ llama.cpp |
| Vision LLM | SmolVLM / Qwen2.5-VL |
| Speech-to-Text | Whisper (រួមទាំងម៉ូដែលដែលបានកែសម្រួលសម្រាប់ភាសាខ្មែរ) |
| Embedding | BAAI/bge-m3 |
| Vector store | ChromaDB (`./chroma_db/`) |
| Visual Index | ColSmolVLM / ColQwen2 (`./visual_index/`) |
| ប្រភេទឯកសារដែលបញ្ចូលបាន | PDF, TXT, MD, DOCX |
| ប្រភេទភ្នាក់ងារ | `smolagents.CodeAgent` |
| UI | Gradio |

## ល្បឿនរំពឹងទុក ({device})

| កិច្ចការ | ពេលវេលា |
|---|---|
| បញ្ចូលឯកសារ | ១០–៦០ វិ |
| ឆ្លើយបែប General / RAG | ១–៥ នាទី |
| ឆ្លើយបែប Vision | ២–៨ នាទី |

> 💡 ម៉ូដែល GGUF (llama.cpp) ត្រូវបានស្កេនដោយស្វ័យប្រវត្តិពីថតដែលអ្នកកំណត់ (ប្រអប់ "📁 ថតម៉ូដែល GGUF" ខាងលើ ឬអថេរបរិស្ថាន `LLAMA_CPP_MODEL_DIR`) ហើយបង្ហាញនៅក្នុងបញ្ជីទម្លាក់ម៉ូដែលដូចគ្នានឹងម៉ូដែល HuggingFace។
"""


def _about_content_en(device: str, version: str) -> str:
    return f"""
### 🔖 Version {version}

A local, bilingual (Khmer/English), multi-modal RAG agent — general chat, document RAG, vision chat, speech-to-text,
AI-driven CSV/Excel data analysis, and knowledge base management — running entirely on your own machine, using either
HuggingFace (transformers) models or local GGUF models via llama.cpp.

---

## Tabs

| Tab | Description |
|---|---|
| 💬 General Chat | Direct LLM conversation — no retrieval |
| 📚 RAG Chat | Retrieves from knowledge base first, then answers |
| 🖼️ Vision Chat | Image understanding with optional text context |
| 🎙️ Speech to Text | Transcribe audio into text using Whisper |
| 📊 Data Analysis | Agent analyzes CSV/Excel, builds charts and a report |
| 📂 Knowledge Base | Upload (PDF / TXT / MD / DOCX) & manage indexed documents |

## Architecture

| Component | Detail |
|---|---|
| LLM | HuggingFace models (Qwen3, Gemma-4) **or** local GGUF models via llama.cpp |
| Vision LLM | SmolVLM / Qwen2.5-VL |
| Speech-to-Text | Whisper (including Khmer-tuned variants) |
| Embedding | BAAI/bge-m3 |
| Vector store | ChromaDB (`./chroma_db/`) |
| Visual Index | ColSmolVLM / ColQwen2 (`./visual_index/`) |
| Supported document types | PDF, TXT, MD, DOCX |
| Agent type | `smolagents.CodeAgent` |
| UI | Gradio |

## Expected speed ({device})

| Task | Time |
|---|---|
| Index a document | 10–60 s |
| General / RAG answer | 1–5 min |
| Vision answer | 2–8 min |

> 💡 GGUF (llama.cpp) models are auto-discovered from whatever folder you configure (the "📁 GGUF Model Folder" box above, or the `LLAMA_CPP_MODEL_DIR` environment variable) and appear in the same model dropdown as the HuggingFace models.
"""

BASE_MODEL_OPTIONS = {
    "🟢 Qwen3-0.6B   (~1.2 GB RAM | fastest)": "Qwen/Qwen3-0.6B",
    "🟡 Qwen3-1.7B   (~3 GB RAM)":             "Qwen/Qwen3-1.7B",
    "🟡 Qwen3-4B     (~7 GB RAM)":             "Qwen/Qwen3-4B",
    "🔵 Gemma-4-E2B  (~4 GB RAM)":             "google/gemma-4-E2B-it",
}

# MODEL_OPTIONS starts as a copy of the base HuggingFace models. Any local
# .gguf models found under LLAMA_CPP_MODEL_DIR are merged in on top of it so
# they appear in the same dropdowns. The folder is user-configurable — via
# the LLAMA_CPP_MODEL_DIR environment variable at startup, or live from the
# "📁 GGUF Model Folder" box in the UI (see rescan_gguf_models() below).
MODEL_OPTIONS = dict(BASE_MODEL_OPTIONS)
MODEL_OPTIONS.update(discover_gguf_models())

DEFAULT_LLM_LABEL = "🟢 Qwen3-0.6B   (~1.2 GB RAM | fastest)"
DEFAULT_LLM_MODEL = MODEL_OPTIONS[DEFAULT_LLM_LABEL]


def rescan_gguf_models(folder_path: Optional[str], lang_key: str = "kh"):
    """Set (or change) the GGUF model folder at runtime and rebuild
    MODEL_OPTIONS, without needing to edit code or restart the app.

    Returns (status_message, dropdown_update) repeated for every model
    dropdown in the UI so they all refresh with the newly discovered
    .gguf models immediately.
    """
    global LLAMA_CPP_MODEL_DIR, MODEL_OPTIONS
    l = LANGUAGES.get(lang_key, LANGUAGES["kh"])
    folder_path = (folder_path or "").strip()
    LLAMA_CPP_MODEL_DIR = folder_path

    MODEL_OPTIONS.clear()
    MODEL_OPTIONS.update(BASE_MODEL_OPTIONS)

    if not folder_path:
        msg = l["gguf_scan_disabled"]
    elif not LLAMA_CPP_AVAILABLE:
        msg = l["gguf_scan_no_backend"]
    else:
        found = discover_gguf_models(folder_path)
        MODEL_OPTIONS.update(found)
        if found:
            msg = l["gguf_scan_found"].format(n=len(found), dir=folder_path)
        else:
            msg = l["gguf_scan_empty"].format(dir=folder_path)

    choices = list(MODEL_OPTIONS.keys())
    default_value = DEFAULT_LLM_LABEL if DEFAULT_LLM_LABEL in choices else (choices[0] if choices else None)
    dd_update = gr.update(choices=choices, value=default_value)
    return msg, dd_update, dd_update, dd_update

VLM_OPTIONS = {
    "🔵 SmolVLM-256M  (~0.5 GB RAM | tiny)":  "HuggingFaceTB/SmolVLM-256M-Instruct",
    "🔵 SmolVLM-500M  (~1 GB RAM | recommended)": "HuggingFaceTB/SmolVLM-500M-Instruct",
    "🟢 Qwen2.5-VL-3B (~6 GB RAM)":           "Qwen/Qwen2.5-VL-3B-Instruct",
}
DEFAULT_VLM_LABEL = "🔵 SmolVLM-500M  (~1 GB RAM | recommended)"
DEFAULT_VLM_MODEL = VLM_OPTIONS[DEFAULT_VLM_LABEL]

VISUAL_RETRIEVER_OPTIONS = {
    "vidore/colsmolvlm-v0.1  (~2 GB | recommended)": "vidore/colsmolvlm-v0.1",
    "vidore/colqwen2-v1.0    (~8 GB | higher accuracy)": "vidore/colqwen2-v1.0",
}
DEFAULT_VISUAL_RETRIEVER = "vidore/colsmolvlm-v0.1"

STT_OPTIONS = {
    "🟢 Whisper-tiny    (~1 GB RAM | fastest)":   "openai/whisper-tiny",
    "🟡 Whisper-base    (~1 GB RAM)":              "openai/whisper-base",
    "🟡 Whisper-small   (~2 GB RAM | recommended)": "openai/whisper-small",
    "🔵 Whisper-large-v3 (~10 GB RAM | best accuracy, multilingual incl. Khmer)": "openai/whisper-large-v3",
    "🇰🇭 Whisper-small — ខ្មែរ (~1 GB RAM | Khmer-tuned)": "seanghay/whisper-small-khmer-v2",
    "🇰🇭 Whisper-large-v3-turbo — ខ្មែរ (~6 GB RAM | best for Khmer)": "metythorn/whisper-large-v3-turbo-mixed-20eps-clean-text-197k",
}
DEFAULT_STT_LABEL = "🟡 Whisper-small   (~2 GB RAM | recommended)"
DEFAULT_STT_MODEL = STT_OPTIONS[DEFAULT_STT_LABEL]

DEFAULT_EMBED_MODEL = "BAAI/bge-m3"
CHROMA_PERSIST_DIR  = "./chroma_db"
VISUAL_INDEX_DIR    = "./visual_index"
CHUNK_SIZE          = 1024
CHUNK_OVERLAP       = 128
TOP_K               = 4
MAX_NEW_TOKENS      = 512

DATA_ANALYSIS_DIR   = "./data_analysis"
DATA_UPLOAD_DIR      = os.path.join(DATA_ANALYSIS_DIR, "uploads")
DATA_OUTPUT_DIR      = os.path.join(DATA_ANALYSIS_DIR, "outputs")
DATA_AGENT_MAX_STEPS = 12

QWEN3_IDS   = {"Qwen/Qwen3-0.6B","Qwen/Qwen3-1.7B","Qwen/Qwen3-4B",
               "Qwen/Qwen3-8B","Qwen/Qwen3-14B","Qwen/Qwen3-32B"}
QWEN36_IDS  = {"Qwen/Qwen3.6-27B","Qwen/Qwen3.6-35B-A3B"}
ALL_QWEN_IDS= QWEN3_IDS | QWEN36_IDS
QWEN_VL_IDS = {"Qwen/Qwen2.5-VL-3B-Instruct","Qwen/Qwen2.5-VL-7B-Instruct"}
SMOL_VLM_IDS= {"HuggingFaceTB/SmolVLM-256M-Instruct","HuggingFaceTB/SmolVLM-500M-Instruct",
               "HuggingFaceTB/SmolVLM2-2.2B-Instruct"}

_embed_model         = None
_chroma_col          = None
_llm                 = None
_llm_model_id        = DEFAULT_LLM_MODEL
_vlm_model           = None
_vlm_processor       = None
_vlm_model_id        = None
_visual_retriever    = None
_visual_retriever_id = None
_llm_lock            = threading.Lock()
_vlm_lock            = threading.Lock()
_stt_pipeline        = None
_stt_model_id        = None
_stt_lock            = threading.Lock()
_data_agent          = None
_data_agent_model_id = None
_data_agent_lock     = threading.Lock()


def _release_model(obj):
    """Explicitly free memory held by a model before dropping the last
    Python reference to it.

    Just setting a global cache variable to None does NOT free GPU/CPU
    memory right away:
      - transformers/torch models: the CUDA caching allocator keeps the
        freed blocks around until torch.cuda.empty_cache() runs (and that
        can only reclaim memory from tensors that have already been
        garbage-collected).
      - llama.cpp (llama-cpp-python) models: the native context / KV-cache
        lives in a C++ object. Its Python wrapper's __del__ will eventually
        close it, but only once the GC actually runs the finalizer — which
        is not guaranteed to happen before the next model tries to load,
        causing an OOM when switching models back-to-back.

    Calling this before overwriting/loading a new model avoids the
    "previous model didn't offload, so the next one fails to load" issue.
    """
    if obj is None:
        return
    try:
        # llama.cpp backend: LlamaCppModel wraps the native llama_cpp.Llama
        # instance as `.llm` — close it explicitly to free the KV-cache /
        # context right now instead of waiting on GC.
        inner = getattr(obj, "llm", None)
        if inner is not None and hasattr(inner, "close"):
            try:
                inner.close()
            except Exception:
                pass

        # transformers / smolagents TransformersModel: move the underlying
        # nn.Module off the GPU before dropping it so CUDA's allocator can
        # actually reclaim the memory once we empty_cache() below.
        model_attr = getattr(obj, "model", None)
        if model_attr is not None and hasattr(model_attr, "to"):
            try:
                model_attr.to("cpu")
            except Exception:
                pass
    except Exception:
        pass

    del obj
    gc.collect()
    try:
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            torch.cuda.ipc_collect()
        elif hasattr(torch.backends, "mps") and torch.backends.mps.is_available():
            torch.mps.empty_cache()
    except Exception:
        pass


def get_embed_model():
    global _embed_model
    if _embed_model is None:
        print(f"[RAG] Loading embedding model '{DEFAULT_EMBED_MODEL}' on {DEVICE.upper()} …")
        from sentence_transformers import SentenceTransformer
        _embed_model = SentenceTransformer(DEFAULT_EMBED_MODEL, device=DEVICE)
    return _embed_model


def encode_texts(texts: list, normalize: bool = True):
    vecs = get_embed_model().encode(texts, normalize_embeddings=normalize,
                                    show_progress_bar=False)
    return vecs.tolist() if hasattr(vecs, "tolist") else vecs


def get_chroma_collection(name: str = "rag_docs"):
    global _chroma_col
    if _chroma_col is None:
        import chromadb
        client      = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
        _chroma_col = client.get_or_create_collection(
            name=name, metadata={"hnsw:space": "cosine"})
        print(f"[RAG] ChromaDB ready — {_chroma_col.count()} chunks indexed.")
    return _chroma_col


def get_llm(model_id: Optional[str] = None):
    global _llm, _llm_model_id
    target = model_id or _llm_model_id

    if _llm is not None and target == _llm_model_id:
        return _llm

    with _llm_lock:
        if _llm is not None and target == _llm_model_id:
            return _llm

        if _llm is not None:
            print(f"[RAG] Unloading previous LLM '{_llm_model_id}' before loading '{target}' …")
            _release_model(_llm)
            _llm = None

        if str(target).lower().endswith(".gguf"):
            if not LLAMA_CPP_AVAILABLE:
                raise RuntimeError(
                    "llama-cpp-python is not installed. Run SETUP.bat to install "
                    "a hardware-matched build, or install it manually with the "
                    "appropriate CUDA/Metal/ROCm build flags for GPU support."
                )
            _llm = LlamaCppModel(
                model_path=target,
                n_ctx=16384,
                # -1 offloads every layer to GPU; on a CPU-only llama-cpp-python
                # build this is harmless (llama.cpp silently ignores it and runs
                # on CPU), but we set 0 explicitly once we know for sure so the
                # load logs are accurate instead of claiming a GPU offload that
                # won't happen.
                n_gpu_layers=-1 if LLAMA_CPP_GPU_AVAILABLE else 0,
                # Flash attention only helps (and is only reliably supported)
                # when the model is actually running on GPU; LlamaCppModel
                # itself also falls back gracefully if a specific model's
                # architecture rejects FA even when the GPU build supports it.
                flash_attn=LLAMA_CPP_GPU_AVAILABLE,
                temperature=0.6,
                top_p=0.95,
                max_new_tokens=MAX_NEW_TOKENS,
            )
        else:
            from smolagents import TransformersModel
            print(f"[RAG] Loading LLM '{target}' on {DEVICE.upper()} …")
            _llm = TransformersModel(
                model_id=target,
                device_map=DEVICE,
                torch_dtype=TORCH_DTYPE,
                max_new_tokens=MAX_NEW_TOKENS,
                temperature=0.6,
                top_p=0.95,
                trust_remote_code=True,
            )
        _llm_model_id = target
        return _llm


def _call_llm(model_id: str, system: str, user: str) -> tuple[str, float]:
    from smolagents.models import ChatMessage
    llm = get_llm(model_id)
    messages = [
        ChatMessage(role="system", content=[{"type": "text", "text": system}]),
        ChatMessage(role="user",   content=[{"type": "text", "text": user}]),
    ]
    t0  = time.time()
    out = llm(messages)
    if hasattr(out, "content"):
        c = out.content
        if isinstance(c, list):
            ans = " ".join(b.get("text", "") if isinstance(b, dict) else str(b) for b in c)
        else:
            ans = str(c)
    else:
        ans = str(out)
    return ans.strip(), time.time() - t0


def get_vlm(model_id: Optional[str] = None):
    global _vlm_model, _vlm_processor, _vlm_model_id

    target = model_id or DEFAULT_VLM_MODEL

    if _vlm_model is not None and target == _vlm_model_id:
        return _vlm_model, _vlm_processor

    if _vlm_model is not None:
        print(f"[VLM] Unloading previous VLM '{_vlm_model_id}' before loading '{target}' …")
        _release_model(_vlm_model)
        _vlm_model = None
        _vlm_processor = None

    print(f"[VLM] Loading '{target}' on {DEVICE.upper()} …")

    from transformers import AutoProcessor

    if target in QWEN_VL_IDS:
        from transformers import Qwen2_5_VLForConditionalGeneration

        _vlm_processor = AutoProcessor.from_pretrained(
            target,
            trust_remote_code=True,
            min_pixels=256 * 28 * 28,
            max_pixels=1280 * 28 * 28,
        )

        _vlm_model = Qwen2_5_VLForConditionalGeneration.from_pretrained(
            target,
            torch_dtype=TORCH_DTYPE,
            device_map=DEVICE,
            trust_remote_code=True,
        )

        _vlm_model._arch = "qwen_vl"

    elif target in SMOL_VLM_IDS:
        try:
            from transformers import AutoModelForImageTextToText
            ModelClass = AutoModelForImageTextToText
            print("[VLM] Using AutoModelForImageTextToText")
        except ImportError:
            try:
                from transformers import AutoModelForVision2Seq
                ModelClass = AutoModelForVision2Seq
                print("[VLM] Using AutoModelForVision2Seq")
            except ImportError:
                from transformers import AutoModel
                ModelClass = AutoModel
                print("[VLM] Falling back to AutoModel")

        _vlm_processor = AutoProcessor.from_pretrained(
            target,
            trust_remote_code=True,
        )

        _vlm_model = ModelClass.from_pretrained(
            target,
            torch_dtype=TORCH_DTYPE,
            device_map=DEVICE,
            trust_remote_code=True,
        )

        _vlm_model._arch = "smolvlm"

    else:
        raise ValueError(f"Unknown VLM: {target}")

    _vlm_model_id = target

    return _vlm_model, _vlm_processor


def vlm_answer(question: str, images: list, context: str = "", model_id: Optional[str] = None) -> str:
    try:
        model, processor = get_vlm(model_id)
        arch = getattr(model, "_arch", "smolvlm")
        system_prompt = "You are a helpful assistant. Answer based on images and context."
        user_text = question + (f"\n\nContext:\n{context}" if context else "")
        if arch == "qwen_vl":
            from qwen_vl_utils import process_vision_info
            content = [{"type": "image", "image": img} for img in images]
            content.append({"type": "text", "text": user_text})
            messages = [{"role": "system", "content": system_prompt},
                        {"role": "user",   "content": content}]
            text_in = processor.apply_chat_template(messages, tokenize=False,
                                                    add_generation_prompt=True)
            img_in, _ = process_vision_info(messages)
            inputs = processor(text=[text_in], images=img_in,
                               padding=True, return_tensors="pt").to(DEVICE)
        else:
            content = [{"type": "image"} for _ in images]
            content.append({"type": "text", "text": user_text})
            messages = [{"role": "system", "content": system_prompt},
                        {"role": "user",   "content": content}]
            text_in = processor.apply_chat_template(messages, add_generation_prompt=True)
            inputs  = processor(text=text_in, images=images or None, return_tensors="pt").to(DEVICE)
        with torch.no_grad():
            out = model.generate(**inputs, max_new_tokens=MAX_NEW_TOKENS)
        trimmed = out[0][inputs["input_ids"].shape[-1]:]
        return processor.decode(trimmed, skip_special_tokens=True)
    except Exception as e:
        import traceback
        return f"❌ Error during VLM generation: {str(e)}\n\n{traceback.format_exc()}"


def get_stt_pipeline(model_id: Optional[str] = None):
    global _stt_pipeline, _stt_model_id
    target = model_id or DEFAULT_STT_MODEL

    if _stt_pipeline is not None and target == _stt_model_id:
        return _stt_pipeline

    with _stt_lock:
        if _stt_pipeline is not None and target == _stt_model_id:
            return _stt_pipeline

        if _stt_pipeline is not None:
            print(f"[STT] Unloading previous STT model '{_stt_model_id}' before loading '{target}' …")
            _release_model(_stt_pipeline)
            _stt_pipeline = None

        print(f"[STT] Loading '{target}' on {DEVICE.upper()} …")
        from transformers import pipeline as hf_pipeline

        # pipeline() wants an integer device index for cuda, "mps", or -1 for cpu
        if DEVICE == "cuda":
            device_arg = 0
        elif DEVICE == "mps":
            device_arg = "mps"
        else:
            device_arg = -1

        _stt_pipeline = hf_pipeline(
            "automatic-speech-recognition",
            model=target,
            torch_dtype=TORCH_DTYPE,
            device=device_arg,
        )
        _stt_model_id = target
        return _stt_pipeline


def transcribe_audio(audio_path: Optional[str], language: Optional[str] = None,
                     model_id: Optional[str] = None) -> str:
    if not audio_path:
        return ""
    try:
        asr = get_stt_pipeline(model_id)
        gen_kwargs = {}
        if language and language != "auto":
            gen_kwargs["language"] = language
            gen_kwargs["task"] = "transcribe"

        # Decode the audio file ourselves via soundfile/librosa (both already
        # in requirements.txt) instead of handing transformers a raw filepath,
        # which requires ffmpeg to be separately installed and on PATH.
        import numpy as np
        target_sr = getattr(getattr(asr, "feature_extractor", None), "sampling_rate", 16000)

        try:
            import soundfile as sf
            audio_array, src_sr = sf.read(audio_path, dtype="float32", always_2d=False)
        except Exception:
            import librosa
            audio_array, src_sr = librosa.load(audio_path, sr=None, mono=False)

        audio_array = np.asarray(audio_array, dtype="float32")
        if audio_array.ndim == 2:
            # Collapse multi-channel audio to mono (channel axis is whichever is smaller)
            channel_axis = 0 if audio_array.shape[0] < audio_array.shape[1] else 1
            audio_array = audio_array.mean(axis=channel_axis).astype("float32")

        if src_sr != target_sr:
            import librosa
            audio_array = librosa.resample(audio_array, orig_sr=src_sr, target_sr=target_sr)

        result = asr({"array": audio_array, "sampling_rate": target_sr}, generate_kwargs=gen_kwargs)
        text = result.get("text", "") if isinstance(result, dict) else str(result)
        return text.strip()
    except Exception as e:
        import traceback
        return f"❌ {e}\n\n{traceback.format_exc()}"


# ──────────────────────────────────────────────────────────────────
# Data Analysis — smolagents CodeAgent w/ pip-install tool
# ──────────────────────────────────────────────────────────────────
from smolagents import CodeAgent, tool


@tool
def install_package(package_name: str) -> str:
    """
    Install a Python package into the current environment using pip.
    Use this whenever a data-analysis step needs a library that is not
    yet installed (e.g. "openpyxl" for reading .xlsx files, "seaborn",
    "scikit-learn", "statsmodels", "plotly", "xlsxwriter").

    Args:
        package_name: The pip package name to install, e.g. "seaborn" or
            "scikit-learn==1.4.0". Pass a single package per call.
    """
    try:
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "--quiet", "--disable-pip-version-check", package_name],
            capture_output=True, text=True, timeout=300,
        )
        if result.returncode == 0:
            return f"✅ Installed '{package_name}' successfully."
        return f"❌ Failed to install '{package_name}':\n{result.stderr[-2000:]}"
    except subprocess.TimeoutExpired:
        return f"❌ Installing '{package_name}' timed out after 300s."
    except Exception as e:
        return f"❌ Error installing '{package_name}': {e}"


def get_data_agent(model_id: Optional[str] = None):
    """Lazily build (or rebuild, if the model changed) the data-analysis CodeAgent."""
    global _data_agent, _data_agent_model_id
    target = model_id or _llm_model_id

    if _data_agent is not None and target == _data_agent_model_id:
        return _data_agent

    with _data_agent_lock:
        if _data_agent is not None and target == _data_agent_model_id:
            return _data_agent

        print(f"[DataAgent] Building CodeAgent on '{target}' …")
        llm = get_llm(target)
        _data_agent = CodeAgent(
            model=llm,
            tools=[install_package],
            additional_authorized_imports=["*"],   # trusted local machine — full stdlib + installed pkgs
            max_steps=DATA_AGENT_MAX_STEPS,
        )
        _data_agent_model_id = target
        return _data_agent


def save_data_files(files) -> list:
    """Copy uploaded CSV/XLSX files into the persistent data-analysis workspace."""
    Path(DATA_UPLOAD_DIR).mkdir(parents=True, exist_ok=True)
    paths = []
    if not files:
        return paths
    if not isinstance(files, list):
        files = [files]
    for f in files:
        src = _get_file_path(f)
        if not src:
            continue
        dest = Path(DATA_UPLOAD_DIR) / Path(src).name
        try:
            shutil.copy(src, dest)
            paths.append(str(dest))
        except Exception:
            pass
    return paths


def run_data_analysis(files, question: str, model_label: str, history: list):
    """Hand uploaded data + the user's question to the CodeAgent and collect its report."""
    history = history or []

    paths = save_data_files(files)
    if not paths:
        history.append({"role": "user", "content": question or "(no file)"})
        history.append({"role": "assistant", "content": "⚠️ Please upload a CSV or XLSX file first."})
        return history, None, None

    question = (question or "").strip() or (
        "Explore this dataset, summarize key statistics and trends, "
        "and generate a short report with at least one chart."
    )
    model_id = MODEL_OPTIONS.get(model_label, DEFAULT_LLM_MODEL)
    history.append({"role": "user", "content": f"📎 {', '.join(Path(p).name for p in paths)}\n\n{question}"})

    try:
        Path(DATA_OUTPUT_DIR).mkdir(parents=True, exist_ok=True)
        agent = get_data_agent(model_id)
        file_list_str = "\n".join(f"- {p}" for p in paths)
        report_path = str(Path(DATA_OUTPUT_DIR) / "report.md")

        task = f"""You are a data analysis assistant working with pandas in a local Python sandbox.

Data file(s) provided by the user:
{file_list_str}

User request: {question}

Instructions:
1. Load the file(s) with pandas (pd.read_csv for .csv, pd.read_excel for .xlsx/.xls —
   if a required package like 'openpyxl' is missing, call the install_package tool with
   its pip name first, then retry the import).
2. Explore the data: shape, columns, dtypes, missing values, basic descriptive statistics.
3. Perform the analysis the user asked for. If a package you need isn't available,
   install it with the install_package tool instead of giving up.
4. Create at least one relevant chart with matplotlib and save each chart as a PNG file
   inside the directory '{DATA_OUTPUT_DIR}' (use plt.savefig(...); do not call plt.show()).
5. Write a concise Markdown report of your findings (headings, bullet points, key numbers)
   and save it to '{report_path}'.
6. As your FINAL ANSWER, return the full Markdown report text.
"""
        t0 = time.time()
        result = agent.run(task)
        elapsed = time.time() - t0

        report_text = str(result)
        response = (
            report_text +
            f"\n\n<hr><sub>⏱ {elapsed:.1f}s | model: <code>{model_id}</code> ({DEVICE.upper()})</sub>"
        )
        history.append({"role": "assistant", "content": response})

        chart_files = sorted(str(p) for p in Path(DATA_OUTPUT_DIR).glob("*.png"))
        report_file = report_path if Path(report_path).exists() else None
        return history, (chart_files or None), report_file

    except Exception as e:
        import traceback
        history.append({"role": "assistant", "content": f"❌ {e}\n\n{traceback.format_exc()}"})
        return history, None, None


def get_visual_retriever():
    global _visual_retriever

    if _visual_retriever is not None:
        return _visual_retriever

    target = DEFAULT_VISUAL_RETRIEVER

    try:
        print(f"[Vision] Loading visual retriever: {target}")

        _visual_retriever = RAGMultiModalModel.from_pretrained(
            target,
            verbose=0,
        )

        print(f"[Vision] Loaded: {target}")

    except ValueError as e:
        # Older Byaldi versions only support ColPali / ColQwen2
        if "only supports ColPali and ColQwen2" not in str(e):
            raise

        fallback = "vidore/colqwen2-v1.0"

        print(f"[Vision] '{target}' is not supported by this version of Byaldi.")
        print(f"[Vision] Falling back to: {fallback}")

        _visual_retriever = RAGMultiModalModel.from_pretrained(
            fallback,
            verbose=0,
        )

    return _visual_retriever


def index_pdf_visual(filepath: str, retriever_id: str) -> str:
    try:
        from byaldi import RAGMultiModalModel
        global _visual_retriever, _visual_retriever_id
        index_path = Path(VISUAL_INDEX_DIR) / "main"
        index_path.parent.mkdir(parents=True, exist_ok=True)
        if _visual_retriever is None or _visual_retriever_id != retriever_id:
            _visual_retriever = RAGMultiModalModel.from_pretrained(retriever_id, verbose=0)
            _visual_retriever_id = retriever_id
        if index_path.exists():
            _visual_retriever.add_to_index(input_item=filepath,
                                           store_collection_with_index=True,
                                           doc_id=int(time.time()))
        else:
            _visual_retriever.index(input_path=filepath, index_name="main",
                                    index_root=VISUAL_INDEX_DIR,
                                    store_collection_with_index=True, overwrite=False)
        return f"✅ Visual index updated: '{Path(filepath).name}'"
    except ImportError:
        return "⚠️ byaldi not installed — skipping visual index."
    except Exception as e:
        return f"❌ {e}"


def visual_retrieve(query: str, top_k: int = 3) -> list:
    retriever = get_visual_retriever()
    if retriever is None:
        return []
    try:
        results = retriever.search(query, k=top_k)
        images  = []
        for r in results:
            if hasattr(r, "base64") and r.base64:
                images.append(Image.open(BytesIO(base64.b64decode(r.base64))).convert("RGB"))
        return images
    except Exception as e:
        print(f"[VIS] {e}")
        return []


def _chunk_text(text: str):
    chunks, start = [], 0
    while start < len(text):
        chunks.append(text[start: start + CHUNK_SIZE])
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks


def _safe_meta(meta: dict, chunk_index: int) -> dict:
    safe = {k: v if isinstance(v, (str, int, float, bool)) else str(v)
            for k, v in meta.items()}
    safe["chunk_index"] = chunk_index
    return safe


def index_texts(texts: list, metadatas: list) -> int:
    col = get_chroma_collection()
    all_chunks, all_metas, all_ids = [], [], []
    for text, meta in zip(texts, metadatas):
        for j, chunk in enumerate(_chunk_text(text)):
            uid = f"{meta.get('source','doc')}__{len(all_chunks)}"
            all_chunks.append(chunk)
            all_metas.append(_safe_meta(meta, j))
            all_ids.append(uid)
    if not all_chunks:
        return 0
    for i in range(0, len(all_chunks), 64):
        vecs = encode_texts(all_chunks[i:i+64])
        col.upsert(embeddings=vecs, documents=all_chunks[i:i+64],
                   metadatas=all_metas[i:i+64], ids=all_ids[i:i+64])
    return len(all_chunks)


def index_pdf_file(filepath: str, visual_retriever_id: str = DEFAULT_VISUAL_RETRIEVER) -> str:
    msgs = []
    try:
        import fitz
        doc = fitz.open(filepath)
        texts, metas = [], []
        for n, page in enumerate(doc):
            t = page.get_text()
            if t.strip():
                texts.append(t)
                metas.append({"source": Path(filepath).name, "page": n+1, "type": "pdf"})
        doc.close()
        msgs.append(f"✅ Text: {index_texts(texts, metas)} chunks")
    except Exception as e:
        msgs.append(f"⚠️ Text failed: {e}")
    msgs.append(f"🖼️ {index_pdf_visual(filepath, visual_retriever_id)}")
    return "\n".join(msgs)


def index_txt_file(filepath: str) -> str:
    try:
        text = Path(filepath).read_text(encoding="utf-8", errors="replace")
        k    = index_texts([text], [{"source": Path(filepath).name, "type": "txt"}])
        return f"✅ {k} chunks from '{Path(filepath).name}'"
    except Exception as e:
        return f"❌ {e}"


def index_docx_file(filepath: str) -> str:
    try:
        import docx  # python-docx
        document = docx.Document(filepath)

        parts = []
        # Paragraph text (skips empty lines)
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table text — tables aren't covered by document.paragraphs
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        if not text.strip():
            return f"⚠️ No extractable text in '{Path(filepath).name}'"

        k = index_texts([text], [{"source": Path(filepath).name, "type": "docx"}])
        return f"✅ {k} chunks from '{Path(filepath).name}'"
    except ImportError:
        return "⚠️ python-docx not installed — run 'pip install python-docx' to index .docx files."
    except Exception as e:
        return f"❌ {e}"


def _get_file_path(f) -> Optional[str]:
    if f is None:          return None
    if isinstance(f, str): return f
    if isinstance(f, dict):return f.get("path") or f.get("name") or f.get("tmp_path")
    if hasattr(f, "path"): return f.path
    if hasattr(f, "name"): return f.name
    return None


def index_uploaded_files(files, visual_retriever_label: str) -> str:
    if not files:
        return "No files uploaded."
    retriever_id = VISUAL_RETRIEVER_OPTIONS.get(visual_retriever_label, DEFAULT_VISUAL_RETRIEVER)
    if not isinstance(files, list):
        files = [files]
    flat = []
    for item in files:
        flat.extend(item) if isinstance(item, list) else flat.append(item)
    msgs = []
    for f in flat:
        path = _get_file_path(f)
        if not path:
            msgs.append(f"⚠️ Cannot resolve path: {f!r}")
            continue
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            msgs.append(index_pdf_file(path, retriever_id))
        elif ext in (".txt", ".md"):
            msgs.append(index_txt_file(path))
        elif ext == ".docx":
            msgs.append(index_docx_file(path))
        else:
            msgs.append(f"⚠️ Unsupported: {ext}")
    return "\n".join(msgs) or "Nothing indexed."


def index_hf_dataset(dataset_name: str, text_col: str, source_col: str = ""):
    try:
        import datasets as ds
        dataset = ds.load_dataset(dataset_name, split="train")
        texts, metas = [], []
        for row in dataset:
            t = row.get(text_col, "")
            if not t:
                continue
            src = row.get(source_col, dataset_name) if source_col else dataset_name
            texts.append(str(t))
            metas.append({"source": str(src), "type": "hf_dataset"})
        k = index_texts(texts, metas)
        return f"✅ {k} chunks from '{dataset_name}'", get_doc_table()
    except Exception as e:
        return f"❌ {e}", get_doc_table()


def get_doc_table() -> list:
    col = get_chroma_collection()
    if col.count() == 0:
        return []
    from collections import defaultdict
    result = col.get(include=["metadatas"])
    agg = defaultdict(lambda: {"type": "", "pages": set(), "chunks": 0})
    for m in result["metadatas"]:
        src = m.get("source", "unknown")
        agg[src]["type"]   = m.get("type", "unknown")
        agg[src]["chunks"] += 1
        if m.get("page"):
            agg[src]["pages"].add(m["page"])
    return [[src, info["type"],
             str(len(info["pages"])) if info["pages"] else "—",
             info["chunks"]]
            for src, info in sorted(agg.items())]


def delete_selected_sources(selected_rows: list, doc_table_data: list) -> tuple:
    if not selected_rows:
        return get_doc_table(), "⚠️ No rows selected."
    col, deleted = get_chroma_collection(), []
    for row_idx in selected_rows:
        if row_idx >= len(doc_table_data):
            continue
        src    = doc_table_data[row_idx][0]
        result = col.get(where={"source": src}, include=["metadatas"])
        if result["ids"]:
            col.delete(ids=result["ids"])
            deleted.append(f"'{src}' ({len(result['ids'])} chunks)")
    msg = ("🗑️ Deleted: " + ", ".join(deleted)) if deleted else "⚠️ Nothing deleted."
    return get_doc_table(), msg


def clear_index() -> tuple:
    global _chroma_col
    import chromadb
    client = chromadb.PersistentClient(path=CHROMA_PERSIST_DIR)
    try:
        client.delete_collection("rag_docs")
    except Exception:
        pass
    _chroma_col = None
    get_chroma_collection()
    return [], "🗑️ All documents cleared."


def get_index_stats(lang_key: str = "en") -> str:
    l = LANGUAGES[lang_key]
    n       = get_chroma_collection().count()
    vis_str = l["err_visual_ready"] if (Path(VISUAL_INDEX_DIR) / "main").exists() else l["err_empty_visual"]
    return l["err_status_bar"].format(n=n, vis=vis_str, dev=DEVICE.upper())


def retrieve_context(query: str) -> tuple[str, list[str]]:
    col = get_chroma_collection()
    if col.count() == 0:
        return "", []
    results = col.query(query_embeddings=encode_texts([query]),
                        n_results=min(TOP_K, col.count()))
    docs  = results["documents"][0]
    metas = results["metadatas"][0]
    if not docs:
        return "", []
    context = "\n\n".join(f"[{m.get('source','?')}]\n{d}"
                          for d, m in zip(docs, metas))
    sources = list({m.get("source", "?") for m in metas})
    return context, sources


import re

def format_llm_response(text: str) -> str:
    m = re.search(r"<think>(.*?)</think>(.*)", text, re.DOTALL)

    if not m:
        return text

    thinking = m.group(1).strip()
    answer = m.group(2).strip()

    return f"""
<details style="
margin-bottom:12px;
border:1px solid #555;
border-radius:8px;
background:#2d2d2d;
padding:10px;
">
<summary style="
cursor:pointer;
font-weight:bold;
color:#ffcc66;
">
🧠 Reasoning (click to expand)
</summary>

<div style="
margin-top:10px;
color:#cfcfcf;
font-family:monospace;
white-space:pre-wrap;
line-height:1.5;
">
{thinking}
</div>

</details>

<div style="
border-left:5px solid #4CAF50;
padding:12px;
background:#1f1f1f;
border-radius:8px;
font-size:16px;
line-height:1.6;
">

<b>💬 Answer</b>

{answer}

</div>
"""

def chat_general(user_message: str, history: list, model_label: str):
    if not user_message.strip():
        return history, ""
    history = history or []
    history.append({"role": "user", "content": user_message})
    model_id = MODEL_OPTIONS.get(model_label, DEFAULT_LLM_MODEL)
    try:
        system = "You are a helpful, friendly assistant."
        ans, elapsed = _call_llm(model_id, system, user_message)
        formatted = format_llm_response(ans)

        response = (
            formatted +
            f"\n\n<hr><sub>⏱ {elapsed:.1f}s | model: <code>{model_id}</code> ({DEVICE.upper()})</sub>"
        )

    except Exception as e:
        import traceback
        response = f"❌ {e}\n\n{traceback.format_exc()}"
    history.append({"role": "assistant", "content": response})
    return history, ""


def chat_rag(user_message: str, history: list, model_label: str):
    if not user_message.strip():
        return history, ""
    history = history or []
    history.append({"role": "user", "content": user_message})
    model_id = MODEL_OPTIONS.get(model_label, DEFAULT_LLM_MODEL)
    try:
        context, sources = retrieve_context(user_message)
        system = (
            "You are a helpful assistant. "
            "Answer the user's question using ONLY the provided context. "
            "If the context does not contain the answer, say so clearly."
        )
        if context:
            user_prompt = f"Context:\n{context}\n\nQuestion: {user_message}"
        else:
            user_prompt = (f"Question: {user_message}\n\n"
                           "(The knowledge base is empty — please index some documents first.)")
        ans, elapsed = _call_llm(model_id, system, user_prompt)
        src_str  = f" | sources: {', '.join(sources)}" if sources else " | no docs indexed"
        formatted = format_llm_response(ans)
        response = (
            formatted +
            f"\n\n<hr><sub>⏱ {elapsed:.1f}s | model: <code>{model_id}</code> ({DEVICE.upper()}){src_str}</sub>"
        )
    except Exception as e:
        import traceback
        response = f"❌ {e}\n\n{traceback.format_exc()}"
    history.append({"role": "assistant", "content": response})
    return history, ""


def chat_vision(user_message: str, uploaded_image, history: list,
                vlm_label: str, use_visual_rag: bool):
    if not user_message.strip() and uploaded_image is None:
        return history, None
    history = history or []
    history.append({"role": "user", "content": user_message or "(image)"})
    vlm_id = VLM_OPTIONS.get(vlm_label, DEFAULT_VLM_MODEL)
    try:
        # Pre-load VLM outside vlm_answer so errors surface cleanly
        get_vlm(vlm_id)
        pil_images = []
        if uploaded_image is not None:
            if isinstance(uploaded_image, Image.Image):
                pil_images.append(uploaded_image)
            elif isinstance(uploaded_image, str):
                pil_images.append(Image.open(uploaded_image).convert("RGB"))
        if use_visual_rag and user_message:
            pil_images.extend(visual_retrieve(user_message, top_k=2))
        context, _ = retrieve_context(user_message) if user_message else ("", [])
        t0  = time.time()
        ans = vlm_answer(user_message, pil_images, context=context, model_id=vlm_id)
        elapsed  = time.time() - t0
        response = (f"{ans}\n\n*⏱ {elapsed:.1f}s | VLM: `{vlm_id}` ({DEVICE.upper()})"
                    f" | images: {len(pil_images)}*")
    except Exception as e:
        import traceback
        response = f"❌ {e}\n\n{traceback.format_exc()}"
    history.append({"role": "assistant", "content": response})
    return history, None


CSS = """
.status-bar   { font-size:0.82rem; color:#888; padding:4px 8px; }
.header-wrap  { display:flex; align-items:baseline; gap:12px; margin-bottom:6px; }
.header-title { font-size:1.6rem; font-weight:700; }
.header-sub   { font-size:0.9rem; color:#aaa; }
.dev-logo     { width:56px; height:56px; border-radius:50%; object-fit:cover;
                box-shadow:0 0 8px rgba(120,80,255,0.6); flex-shrink:0; margin-right:4px; }
.beta-badge   { display:inline-block; font-size:0.68rem; font-weight:700; letter-spacing:0.5px;
                color:#1a1a1a; background:#ffcc66; border-radius:999px; padding:2px 9px;
                margin-left:6px; vertical-align:middle; }
"""

def refresh_system_ui():
    status = HardwareManager.get_system_status()
    return (
        "🔄 Status Refreshed",
        status["current_device"].upper(),
        status["cuda_version"] or "N/A",
        "✅ Yes" if status["torch_cuda_available"] else "❌ No",
        "" # Clear logs
    )

def fix_system_ui_generator():
    status = HardwareManager.get_system_status()
    yield (
        "🚀 Starting fix...",
        status["current_device"].upper(),
        status["cuda_version"] or "N/A",
        "✅ Yes" if status["torch_cuda_available"] else "❌ No",
        "Starting..."
    )

    for msg, dev, log in HardwareManager.fix_environment_gen():
        yield (
            msg,
            dev.upper(),
            status["cuda_version"] or "N/A",
            "✅ Yes" if status["torch_cuda_available"] else "❌ No",
            log
        )


def build_ui():
    # Start with Khmer as default
    L = LANGUAGES["kh"]

    with gr.Blocks(title="🔍 RAG Agent") as demo:
        lang_state = gr.State("kh")

        # ── Header ────────────────────────────────────────────────
        with gr.Row():
            with gr.Column(scale=8):
                with gr.Row():
                    if DEVELOPER_LOGO_B64:
                        with gr.Column(scale=0, min_width=64):
                            gr.HTML(
                                f'<img src="data:image/jpeg;base64,{DEVELOPER_LOGO_B64}" '
                                f'class="dev-logo" alt="{DEVELOPER_NAME} logo" />'
                            )
                    with gr.Column():
                        header_title = gr.HTML(
                            f'<div class="header-wrap"><span class="header-title">{L["title"]}</span>'
                            f'<span class="beta-badge">BETA</span></div>'
                        )
                        header_sub = gr.HTML(
                            f'<div class="header-wrap"><span class="header-sub">'
                            f'{L["subtitle"].format(device=DEVICE.upper(), version=APP_VERSION)}</span></div>'
                        )
            with gr.Column(scale=2, variant="panel"):
                lang_dropdown = gr.Dropdown(
                    choices=["Khmer", "English"], value="Khmer",
                    label="🌐 Language", scale=1
                )

        # ── GGUF model folder (optional, user-configurable) ─────────
        # No path is hardcoded — leave blank to skip GGUF entirely, or
        # point it at any folder of .gguf files and click Scan. This can
        # also be pre-set via the LLAMA_CPP_MODEL_DIR environment variable.
        with gr.Row():
            gguf_dir_tb = gr.Textbox(
                value=LLAMA_CPP_MODEL_DIR,
                placeholder=L["gguf_dir_placeholder"],
                label=L["label_gguf_dir"], scale=8,
            )
            scan_gguf_btn = gr.Button(L["btn_scan_gguf"], scale=1)
        gguf_scan_status = gr.Textbox(show_label=False, interactive=False)

        # ── Tabs ──────────────────────────────────────────────────
        with gr.Tabs():

            # ── Tab 1: General Chat ───────────────────────────────
            with gr.Tab(L["tab_general"]) as tab_gen:
                gen_desc    = gr.Markdown(L["tab_general_desc"])
                bot_gen     = gr.Chatbot(height=440)
                with gr.Row():
                    msg_gen  = gr.Textbox(placeholder=L["placeholder_gen"], show_label=False, scale=8)
                    send_gen = gr.Button(L["btn_send"], variant="primary", scale=1)
                with gr.Row():
                    model_dd_gen   = gr.Dropdown(choices=list(MODEL_OPTIONS.keys()), value=DEFAULT_LLM_LABEL, label=L["label_llm"], scale=6)
                    reload_gen     = gr.Button(L["btn_load"], size="sm", scale=2)
                    reload_gen_out = gr.Textbox(show_label=False, interactive=False, scale=4)
                clear_gen = gr.Button(L["btn_clear"], size="sm")

            # ── Tab 2: RAG Chat ───────────────────────────────────
            with gr.Tab(L["tab_rag"]) as tab_rag:
                rag_desc    = gr.Markdown(L["tab_rag_desc"])
                bot_rag     = gr.Chatbot(height=440)
                with gr.Row():
                    msg_rag  = gr.Textbox(placeholder=L["placeholder_rag"], show_label=False, scale=8)
                    send_rag = gr.Button(L["btn_send"], variant="primary", scale=1)
                with gr.Row():
                    model_dd_rag   = gr.Dropdown(choices=list(MODEL_OPTIONS.keys()), value=DEFAULT_LLM_LABEL, label=L["label_llm"], scale=6)
                    reload_rag     = gr.Button(L["btn_load"], size="sm", scale=2)
                    reload_rag_out = gr.Textbox(show_label=False, interactive=False, scale=4)
                clear_rag = gr.Button(L["btn_clear"], size="sm")

            # ── Tab 3: Vision Chat ────────────────────────────────
            with gr.Tab(L["tab_vision"]) as tab_vis:
                vis_desc    = gr.Markdown(L["tab_vision_desc"])
                bot_vis     = gr.Chatbot(height=400)
                with gr.Row():
                    msg_vis    = gr.Textbox(placeholder=L["placeholder_vis"], show_label=False, scale=6)
                    img_upload = gr.Image(type="pil", sources=["upload", "clipboard"], scale=2)
                    send_vis   = gr.Button(L["btn_send"], variant="primary", scale=1)
                with gr.Row():
                    vlm_dd       = gr.Dropdown(choices=list(VLM_OPTIONS.keys()), value=DEFAULT_VLM_LABEL, label=L["label_vlm"], scale=5)
                    vis_rag_chk  = gr.Checkbox(label=L["label_vis_rag"], value=True, scale=2)
                    load_vlm_btn = gr.Button(L["btn_load"], size="sm", scale=2)
                    load_vlm_out = gr.Textbox(show_label=False, interactive=False, scale=3)
                clear_vis = gr.Button(L["btn_clear"], size="sm")

            # ── Tab 4: Speech to Text ─────────────────────────────
            with gr.Tab(L["tab_stt"]) as tab_stt:
                stt_desc  = gr.Markdown(L["tab_stt_desc"])
                stt_audio = gr.Audio(label=L["stt_audio_label"], sources=["microphone", "upload"], type="filepath")
                with gr.Row():
                    stt_dd      = gr.Dropdown(choices=list(STT_OPTIONS.keys()), value=DEFAULT_STT_LABEL, label=L["label_stt"], scale=5)
                    stt_lang_dd = gr.Dropdown(
                        choices=[("Auto-detect", "auto"), ("English", "english"), ("Khmer", "khmer"),
                                 ("French", "french"), ("Chinese", "chinese"), ("Japanese", "japanese")],
                        value="auto", label=L["label_stt_lang"], scale=3,
                    )
                    load_stt_btn = gr.Button(L["btn_load"], size="sm", scale=2)
                    load_stt_out = gr.Textbox(show_label=False, interactive=False, scale=3)
                stt_hint = gr.Markdown(L["stt_khmer_hint"])
                transcribe_btn = gr.Button(L["btn_transcribe"], variant="primary")
                stt_output     = gr.Textbox(label=L["label_res"], lines=8, interactive=True)

            # ── Tab 5: Data Analysis ──────────────────────────────
            with gr.Tab(L["tab_data"]) as tab_data:
                data_desc    = gr.Markdown(L["tab_data_desc"])
                data_file_up = gr.File(label=L["data_file_label"], file_types=[".csv", ".xlsx", ".xls"], file_count="multiple")
                bot_data     = gr.Chatbot(height=420)
                with gr.Row():
                    msg_data  = gr.Textbox(placeholder=L["placeholder_data"], show_label=False, scale=8)
                    send_data = gr.Button(L["btn_send"], variant="primary", scale=1)
                with gr.Row():
                    model_dd_data  = gr.Dropdown(choices=list(MODEL_OPTIONS.keys()), value=DEFAULT_LLM_LABEL, label=L["label_llm"], scale=6)
                    reset_data_btn = gr.Button(L["btn_reset_agent"], size="sm", scale=2)
                    reset_data_out = gr.Textbox(show_label=False, interactive=False, scale=4)
                data_gallery     = gr.Gallery(label=L["label_charts"], columns=3, height=280)
                data_report_file = gr.File(label=L["label_report_file"], interactive=False)
                clear_data = gr.Button(L["btn_clear"], size="sm")

            # ── Tab 6: Knowledge Base ─────────────────────────────
            with gr.Tab(L["tab_kb"]) as tab_kb:
                with gr.Accordion(L["accordion_add"], open=True) as acc_add:
                    file_up    = gr.File(label=L["file_label"], file_types=[".pdf",".txt",".md",".docx"], file_count="multiple")
                    vis_ret_dd = gr.Dropdown(choices=list(VISUAL_RETRIEVER_OPTIONS.keys()), value=list(VISUAL_RETRIEVER_OPTIONS.keys())[0], label=L["label_vis_ret"])
                    up_btn     = gr.Button(L["btn_index"], variant="primary")
                    up_msg     = gr.Textbox(label=L["label_res"], interactive=False, lines=4)

                kb_header = gr.Markdown(L["header_docs"])
                doc_table = gr.Dataframe(
                    headers=L["doc_table_headers"],
                    datatype=["str","str","str","number"],
                    value=get_doc_table,
                    interactive=True, wrap=True,
                )
                with gr.Row():
                    refresh_btn    = gr.Button(L["btn_refresh"], size="sm", scale=2)
                    delete_sel_btn = gr.Button(L["btn_delete"], variant="stop", size="sm", scale=2)
                    clear_all_btn  = gr.Button(L["btn_clear_all"], variant="stop", size="sm", scale=2)
                action_msg         = gr.Textbox(label="", interactive=False, lines=1)
                selected_rows_state = gr.State([])

            # ── Tab 7: About ──────────────────────────────────────
            with gr.Tab(L["tab_about"]) as tab_about:
                about_md_kh = gr.Markdown(_about_content_kh(DEVICE.upper(), APP_VERSION))
                gr.Markdown("---")
                about_md_en = gr.Markdown(_about_content_en(DEVICE.upper(), APP_VERSION))

        status_bar = gr.Textbox(
            value=get_index_stats("kh"), interactive=False,
            show_label=False, elem_classes=["status-bar"]
        )

        # ── Event handlers ────────────────────────────────────────

        # GGUF model folder — rescan updates every model dropdown at once
        scan_gguf_btn.click(
            rescan_gguf_models,
            [gguf_dir_tb, lang_state],
            [gguf_scan_status, model_dd_gen, model_dd_rag, model_dd_data],
        )
        gguf_dir_tb.submit(
            rescan_gguf_models,
            [gguf_dir_tb, lang_state],
            [gguf_scan_status, model_dd_gen, model_dd_rag, model_dd_data],
        )


        # General Chat
        def reload_gen_fn(label):
            global _llm, _data_agent, _data_agent_model_id
            # Free the current model's memory now, rather than just dropping
            # the reference — otherwise the old model can still be occupying
            # VRAM when we try to load the new one right below.
            _release_model(_llm)
            _llm = None
            # The data-analysis agent holds its own reference to this same
            # LLM instance — clear it too so it doesn't keep the old model
            # (or its now-stale weights) alive, and rebuilds against the
            # newly loaded one on next use.
            _data_agent = None
            _data_agent_model_id = None
            mid = MODEL_OPTIONS.get(label, DEFAULT_LLM_MODEL)
            try:
                get_llm(mid); return f"✅ '{mid}' loaded"
            except Exception as e:
                return f"❌ {e}"

        msg_gen.submit(chat_general,  [msg_gen, bot_gen, model_dd_gen], [bot_gen, msg_gen])
        send_gen.click(chat_general,  [msg_gen, bot_gen, model_dd_gen], [bot_gen, msg_gen])
        clear_gen.click(lambda: ([], ""), outputs=[bot_gen, msg_gen])
        reload_gen.click(reload_gen_fn, [model_dd_gen], [reload_gen_out])

        # RAG Chat
        def reload_rag_fn(label):
            global _llm, _data_agent, _data_agent_model_id
            _release_model(_llm)
            _llm = None
            _data_agent = None
            _data_agent_model_id = None
            mid = MODEL_OPTIONS.get(label, DEFAULT_LLM_MODEL)
            try:
                get_llm(mid); return f"✅ '{mid}' loaded"
            except Exception as e:
                return f"❌ {e}"

        msg_rag.submit(chat_rag,  [msg_rag, bot_rag, model_dd_rag], [bot_rag, msg_rag])
        send_rag.click(chat_rag,  [msg_rag, bot_rag, model_dd_rag], [bot_rag, msg_rag])
        clear_rag.click(lambda: ([], ""), outputs=[bot_rag, msg_rag])
        reload_rag.click(reload_rag_fn, [model_dd_rag], [reload_rag_out])

        # Vision Chat
        def load_vlm_fn(label):
            global _vlm_model, _vlm_processor
            _release_model(_vlm_model)
            _vlm_model = _vlm_processor = None
            mid = VLM_OPTIONS.get(label, DEFAULT_VLM_MODEL)
            try:
                get_vlm(mid); return f"✅ '{mid}' loaded"
            except Exception as e:
                return f"❌ {e}"

        send_vis.click(chat_vision,  [msg_vis, img_upload, bot_vis, vlm_dd, vis_rag_chk], [bot_vis, img_upload])
        msg_vis.submit(chat_vision,  [msg_vis, img_upload, bot_vis, vlm_dd, vis_rag_chk], [bot_vis, img_upload])
        clear_vis.click(lambda: ([], None), outputs=[bot_vis, img_upload])
        load_vlm_btn.click(load_vlm_fn, [vlm_dd], [load_vlm_out])

        # Speech to Text
        def load_stt_fn(label):
            global _stt_pipeline, _stt_model_id
            _release_model(_stt_pipeline)
            _stt_pipeline = None; _stt_model_id = None
            mid = STT_OPTIONS.get(label, DEFAULT_STT_MODEL)
            try:
                get_stt_pipeline(mid); return f"✅ '{mid}' loaded"
            except Exception as e:
                return f"❌ {e}"

        def do_transcribe(audio_path, stt_label, lang_choice):
            mid = STT_OPTIONS.get(stt_label, DEFAULT_STT_MODEL)
            return transcribe_audio(audio_path, language=lang_choice, model_id=mid)

        transcribe_btn.click(do_transcribe, [stt_audio, stt_dd, stt_lang_dd], [stt_output])
        load_stt_btn.click(load_stt_fn, [stt_dd], [load_stt_out])

        # Data Analysis
        def reset_data_agent_fn():
            global _data_agent, _data_agent_model_id
            # Just drop the CodeAgent wrapper here — it references the shared
            # _llm managed by get_llm()/reload_gen_fn()/reload_rag_fn(), whose
            # memory is released there when the underlying model actually
            # changes. Rebuilding is cheap since it reuses get_llm(target).
            _data_agent = None
            _data_agent_model_id = None
            return "✅ Agent reset — will rebuild on next run."

        def do_data_analysis(files, question, model_label, history):
            history, gallery, report_file = run_data_analysis(files, question, model_label, history)
            return history, "", gallery, report_file

        send_data.click(do_data_analysis, [data_file_up, msg_data, model_dd_data, bot_data],
                        [bot_data, msg_data, data_gallery, data_report_file])
        msg_data.submit(do_data_analysis, [data_file_up, msg_data, model_dd_data, bot_data],
                        [bot_data, msg_data, data_gallery, data_report_file])
        clear_data.click(lambda: ([], None, None), outputs=[bot_data, data_gallery, data_report_file])
        reset_data_btn.click(reset_data_agent_fn, outputs=[reset_data_out])

        # Knowledge Base
        def on_select(evt: gr.SelectData, current):
            row = evt.index[0]
            if row in current: current.remove(row)
            else:              current.append(row)
            return current

        def do_upload(files, vis_ret_label):
            import traceback
            try:
                return index_uploaded_files(files, vis_ret_label), get_doc_table()
            except Exception as e:
                return f"❌ {traceback.format_exc()}", get_doc_table()

        def do_delete(selected, table_data):
            rows = table_data if isinstance(table_data, list) else table_data.values.tolist()
            new_table, msg = delete_selected_sources(selected, rows)
            return new_table, msg, []

        def do_clear():
            table, msg = clear_index()
            return table, msg, []

        doc_table.select(on_select,        [selected_rows_state], [selected_rows_state])
        up_btn.click(do_upload,            [file_up, vis_ret_dd], [up_msg, doc_table])
        refresh_btn.click(get_doc_table,   outputs=[doc_table])
        delete_sel_btn.click(do_delete,    [selected_rows_state, doc_table], [doc_table, action_msg, selected_rows_state])
        clear_all_btn.click(do_clear,      outputs=[doc_table, action_msg, selected_rows_state])

        # ── Language switcher ─────────────────────────────────────
        def switch_lang(lang_name):
            lk = "kh" if lang_name == "Khmer" else "en"
            l  = LANGUAGES[lk]
            return (
                lk,
                # header
                f'<div class="header-wrap"><span class="header-title">{l["title"]}</span>'
                f'<span class="beta-badge">BETA</span></div>',
                f'<div class="header-wrap"><span class="header-sub">{l["subtitle"].format(device=DEVICE.upper(), version=APP_VERSION)}</span></div>',
                # GGUF model folder
                gr.update(label=l["label_gguf_dir"], placeholder=l["gguf_dir_placeholder"]),
                gr.update(value=l["btn_scan_gguf"]),
                # General Chat
                gr.update(value=l["tab_general_desc"]),
                gr.update(placeholder=l["placeholder_gen"]),
                gr.update(value=l["btn_send"]),
                gr.update(label=l["label_llm"]),
                gr.update(value=l["btn_load"]),
                gr.update(value=l["btn_clear"]),
                # RAG Chat
                gr.update(value=l["tab_rag_desc"]),
                gr.update(placeholder=l["placeholder_rag"]),
                gr.update(value=l["btn_send"]),
                gr.update(label=l["label_llm"]),
                gr.update(value=l["btn_load"]),
                gr.update(value=l["btn_clear"]),
                # Vision Chat
                gr.update(value=l["tab_vision_desc"]),
                gr.update(placeholder=l["placeholder_vis"]),
                gr.update(value=l["btn_send"]),
                gr.update(label=l["label_vlm"]),
                gr.update(label=l["label_vis_rag"]),
                gr.update(value=l["btn_load"]),
                gr.update(value=l["btn_clear"]),
                # STT
                gr.update(value=l["tab_stt_desc"]),
                gr.update(label=l["stt_audio_label"]),
                gr.update(label=l["label_stt"]),
                gr.update(label=l["label_stt_lang"]),
                gr.update(value=l["btn_load"]),
                gr.update(value=l["stt_khmer_hint"]),
                gr.update(value=l["btn_transcribe"]),
                gr.update(label=l["label_res"]),
                # Data Analysis
                gr.update(value=l["tab_data_desc"]),
                gr.update(label=l["data_file_label"]),
                gr.update(placeholder=l["placeholder_data"]),
                gr.update(value=l["btn_send"]),
                gr.update(label=l["label_llm"]),
                gr.update(value=l["btn_reset_agent"]),
                gr.update(label=l["label_charts"]),
                gr.update(label=l["label_report_file"]),
                gr.update(value=l["btn_clear"]),
                # Knowledge Base
                gr.update(label=l["file_label"]),
                gr.update(label=l["label_vis_ret"]),
                gr.update(value=l["btn_index"]),
                gr.update(label=l["label_res"]),
                gr.update(value=l["header_docs"]),
                gr.update(value=l["btn_refresh"]),
                gr.update(value=l["btn_delete"]),
                gr.update(value=l["btn_clear_all"]),
                # Status bar
                get_index_stats(lk),
            )

        _lang_outputs = [
            lang_state, header_title, header_sub,
            gguf_dir_tb, scan_gguf_btn,
            # General Chat
            gen_desc, msg_gen, send_gen, model_dd_gen, reload_gen, clear_gen,
            # RAG Chat
            rag_desc, msg_rag, send_rag, model_dd_rag, reload_rag, clear_rag,
            # Vision Chat
            vis_desc, msg_vis, send_vis, vlm_dd, vis_rag_chk, load_vlm_btn, clear_vis,
            # STT
            stt_desc, stt_audio, stt_dd, stt_lang_dd, load_stt_btn, stt_hint, transcribe_btn, stt_output,
            # Data Analysis
            data_desc, data_file_up, msg_data, send_data, model_dd_data, reset_data_btn,
            data_gallery, data_report_file, clear_data,
            # Knowledge Base
            file_up, vis_ret_dd, up_btn, up_msg,
            kb_header, refresh_btn, delete_sel_btn, clear_all_btn,
            # Status bar
            status_bar,
        ]

        lang_dropdown.change(switch_lang, [lang_dropdown], _lang_outputs)

        # Initialise to Khmer on load
        demo.load(lambda: switch_lang("Khmer"), outputs=_lang_outputs)

        return demo
if __name__ == "__main__":
    print("[RAG] Pre-loading embeddings and ChromaDB …")
    get_embed_model()
    get_chroma_collection()
    demo = build_ui()
    demo.launch(
        server_name="0.0.0.0",
        server_port=7861,
        share=False,
        inbrowser=True,
        theme=gr.themes.Soft(primary_hue="violet"),
        css=CSS,
    )
