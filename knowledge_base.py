"""
knowledge_base.py — Document indexing (PDF/TXT/MD/DOCX + HuggingFace
datasets) into ChromaDB, the optional visual (image-based PDF) index via
byaldi, and retrieval helpers used by the chat tabs.
"""

import base64
import time
from collections import defaultdict
from io import BytesIO
from pathlib import Path
from typing import Optional

from PIL import Image
from byaldi import RAGMultiModalModel

import model_registry as mr
import models
from hardware import DEVICE
from i18n import LANGUAGES

_visual_retriever    = None
_visual_retriever_id = None


def get_file_path(f) -> Optional[str]:
    if f is None:          return None
    if isinstance(f, str): return f
    if isinstance(f, dict):return f.get("path") or f.get("name") or f.get("tmp_path")
    if hasattr(f, "path"): return f.path
    if hasattr(f, "name"): return f.name
    return None


def _chunk_text(text: str):
    chunks, start = [], 0
    while start < len(text):
        chunks.append(text[start: start + mr.CHUNK_SIZE])
        start += mr.CHUNK_SIZE - mr.CHUNK_OVERLAP
    return chunks


def _safe_meta(meta: dict, chunk_index: int) -> dict:
    safe = {k: v if isinstance(v, (str, int, float, bool)) else str(v)
            for k, v in meta.items()}
    safe["chunk_index"] = chunk_index
    return safe


def index_texts(texts: list, metadatas: list) -> int:
    col = models.get_chroma_collection()
    all_chunks, all_metas, all_ids = [], [], []
    for text, meta in zip(texts, metadatas):
        for j, chunk in enumerate(_chunk_text(text)):
            uid = f"{meta.get('source','doc')}__{len(all_chunks)}"
            all_chunks.append(chunk)
            all_metas.append(_safe_meta(meta, j))
            all_ids.append(uid)
    if not all_chunks:
        return 0
    for i in range(0, len(all_chunks), 64):
        vecs = models.encode_texts(all_chunks[i:i+64])
        col.upsert(embeddings=vecs, documents=all_chunks[i:i+64],
                   metadatas=all_metas[i:i+64], ids=all_ids[i:i+64])
    return len(all_chunks)


def get_visual_retriever():
    global _visual_retriever

    if _visual_retriever is not None:
        return _visual_retriever

    target = mr.DEFAULT_VISUAL_RETRIEVER

    try:
        print(f"[Vision] Loading visual retriever: {target}")

        _visual_retriever = RAGMultiModalModel.from_pretrained(
            target,
            verbose=0,
        )

        print(f"[Vision] Loaded: {target}")

    except ValueError as e:
        # Older Byaldi versions only support ColPali / ColQwen2
        if "only supports ColPali and ColQwen2" not in str(e):
            raise

        fallback = "vidore/colqwen2-v1.0"

        print(f"[Vision] '{target}' is not supported by this version of Byaldi.")
        print(f"[Vision] Falling back to: {fallback}")

        _visual_retriever = RAGMultiModalModel.from_pretrained(
            fallback,
            verbose=0,
        )

    return _visual_retriever


def index_pdf_visual(filepath: str, retriever_id: str) -> str:
    try:
        global _visual_retriever, _visual_retriever_id
        index_path = Path(mr.VISUAL_INDEX_DIR) / "main"
        index_path.parent.mkdir(parents=True, exist_ok=True)
        if _visual_retriever is None or _visual_retriever_id != retriever_id:
            _visual_retriever = RAGMultiModalModel.from_pretrained(retriever_id, verbose=0)
            _visual_retriever_id = retriever_id
        if index_path.exists():
            _visual_retriever.add_to_index(input_item=filepath,
                                           store_collection_with_index=True,
                                           doc_id=int(time.time()))
        else:
            _visual_retriever.index(input_path=filepath, index_name="main",
                                    index_root=mr.VISUAL_INDEX_DIR,
                                    store_collection_with_index=True, overwrite=False)
        return f"✅ Visual index updated: '{Path(filepath).name}'"
    except ImportError:
        return "⚠️ byaldi not installed — skipping visual index."
    except Exception as e:
        return f"❌ {e}"


def visual_retrieve(query: str, top_k: int = 3) -> list:
    # Nothing has been indexed into the visual (image-based PDF) index yet —
    # skip entirely instead of loading the ~2-8 GB retriever model just to
    # hit "No passages provided". This is the normal state until the user
    # indexes at least one PDF from the Knowledge Base tab.
    if not (Path(mr.VISUAL_INDEX_DIR) / "main").exists():
        return []
    retriever = get_visual_retriever()
    if retriever is None:
        return []
    try:
        results = retriever.search(query, k=top_k)
        images  = []
        for r in results:
            if hasattr(r, "base64") and r.base64:
                images.append(Image.open(BytesIO(base64.b64decode(r.base64))).convert("RGB"))
        return images
    except Exception as e:
        print(f"[VIS] {e}")
        return []


def unload_visual_retriever_fn(lang_key: str = "kh") -> str:
    global _visual_retriever, _visual_retriever_id
    l = LANGUAGES.get(lang_key, LANGUAGES["kh"])
    if _visual_retriever is None:
        return l["btn_unload_none"]
    mid = _visual_retriever_id or mr.DEFAULT_VISUAL_RETRIEVER
    models._release_model(_visual_retriever)
    _visual_retriever = None
    _visual_retriever_id = None
    return l["msg_unloaded"].format(model=mid)


def index_pdf_file(filepath: str, visual_retriever_id: str = mr.DEFAULT_VISUAL_RETRIEVER) -> str:
    msgs = []
    try:
        import fitz
        doc = fitz.open(filepath)
        texts, metas = [], []
        for n, page in enumerate(doc):
            t = page.get_text()
            if t.strip():
                texts.append(t)
                metas.append({"source": Path(filepath).name, "page": n+1, "type": "pdf"})
        doc.close()
        msgs.append(f"✅ Text: {index_texts(texts, metas)} chunks")
    except Exception as e:
        msgs.append(f"⚠️ Text failed: {e}")
    msgs.append(f"🖼️ {index_pdf_visual(filepath, visual_retriever_id)}")
    return "\n".join(msgs)


def index_txt_file(filepath: str) -> str:
    try:
        text = Path(filepath).read_text(encoding="utf-8", errors="replace")
        k    = index_texts([text], [{"source": Path(filepath).name, "type": "txt"}])
        return f"✅ {k} chunks from '{Path(filepath).name}'"
    except Exception as e:
        return f"❌ {e}"


def index_docx_file(filepath: str) -> str:
    try:
        import docx  # python-docx
        document = docx.Document(filepath)

        parts = []
        # Paragraph text (skips empty lines)
        for para in document.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # Table text — tables aren't covered by document.paragraphs
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        text = "\n".join(parts)
        if not text.strip():
            return f"⚠️ No extractable text in '{Path(filepath).name}'"

        k = index_texts([text], [{"source": Path(filepath).name, "type": "docx"}])
        return f"✅ {k} chunks from '{Path(filepath).name}'"
    except ImportError:
        return "⚠️ python-docx not installed — run 'pip install python-docx' to index .docx files."
    except Exception as e:
        return f"❌ {e}"


def index_uploaded_files(files, visual_retriever_label: str) -> str:
    if not files:
        return "No files uploaded."
    retriever_id = mr.VISUAL_RETRIEVER_OPTIONS.get(visual_retriever_label, mr.DEFAULT_VISUAL_RETRIEVER)
    if not isinstance(files, list):
        files = [files]
    flat = []
    for item in files:
        flat.extend(item) if isinstance(item, list) else flat.append(item)
    msgs = []
    for f in flat:
        path = get_file_path(f)
        if not path:
            msgs.append(f"⚠️ Cannot resolve path: {f!r}")
            continue
        ext = Path(path).suffix.lower()
        if ext == ".pdf":
            msgs.append(index_pdf_file(path, retriever_id))
        elif ext in (".txt", ".md"):
            msgs.append(index_txt_file(path))
        elif ext == ".docx":
            msgs.append(index_docx_file(path))
        else:
            msgs.append(f"⚠️ Unsupported: {ext}")
    return "\n".join(msgs) or "Nothing indexed."


def index_hf_dataset(dataset_name: str, text_col: str, source_col: str = ""):
    try:
        import datasets as ds
        dataset = ds.load_dataset(dataset_name, split="train")
        texts, metas = [], []
        for row in dataset:
            t = row.get(text_col, "")
            if not t:
                continue
            src = row.get(source_col, dataset_name) if source_col else dataset_name
            texts.append(str(t))
            metas.append({"source": str(src), "type": "hf_dataset"})
        k = index_texts(texts, metas)
        return f"✅ {k} chunks from '{dataset_name}'", get_doc_table()
    except Exception as e:
        return f"❌ {e}", get_doc_table()


def get_doc_table() -> list:
    col = models.get_chroma_collection()
    if col.count() == 0:
        return []
    result = col.get(include=["metadatas"])
    agg = defaultdict(lambda: {"type": "", "pages": set(), "chunks": 0})
    for m in result["metadatas"]:
        src = m.get("source", "unknown")
        agg[src]["type"]   = m.get("type", "unknown")
        agg[src]["chunks"] += 1
        if m.get("page"):
            agg[src]["pages"].add(m["page"])
    return [[src, info["type"],
             str(len(info["pages"])) if info["pages"] else "—",
             info["chunks"]]
            for src, info in sorted(agg.items())]


def delete_selected_sources(selected_rows: list, doc_table_data: list) -> tuple:
    if not selected_rows:
        return get_doc_table(), "⚠️ No rows selected."
    col, deleted = models.get_chroma_collection(), []
    for row_idx in selected_rows:
        if row_idx >= len(doc_table_data):
            continue
        src    = doc_table_data[row_idx][0]
        result = col.get(where={"source": src}, include=["metadatas"])
        if result["ids"]:
            col.delete(ids=result["ids"])
            deleted.append(f"'{src}' ({len(result['ids'])} chunks)")
    msg = ("🗑️ Deleted: " + ", ".join(deleted)) if deleted else "⚠️ Nothing deleted."
    return get_doc_table(), msg


def clear_index() -> tuple:
    import chromadb
    client = chromadb.PersistentClient(path=mr.CHROMA_PERSIST_DIR)
    try:
        client.delete_collection("rag_docs")
    except Exception:
        pass
    models.reset_chroma_collection()
    models.get_chroma_collection()
    return [], "🗑️ All documents cleared."


def get_index_stats(lang_key: str = "en") -> str:
    l = LANGUAGES[lang_key]
    n       = models.get_chroma_collection().count()
    vis_str = l["err_visual_ready"] if (Path(mr.VISUAL_INDEX_DIR) / "main").exists() else l["err_empty_visual"]
    return l["err_status_bar"].format(n=n, vis=vis_str, dev=DEVICE.upper())


def retrieve_context(query: str) -> tuple[str, list[str]]:
    col = models.get_chroma_collection()
    if col.count() == 0:
        return "", []
    results = col.query(query_embeddings=models.encode_texts([query]),
                        n_results=min(mr.TOP_K, col.count()))
    docs  = results["documents"][0]
    metas = results["metadatas"][0]
    if not docs:
        return "", []
    context = "\n\n".join(f"[{m.get('source','?')}]\n{d}"
                          for d, m in zip(docs, metas))
    sources = list({m.get("source", "?") for m in metas})
    return context, sources
